import base64
import datetime
import json
import mimetypes
import os
import traceback
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib import messages
from django.core.exceptions import ValidationError, ObjectDoesNotExist
from django.db import DataError, connection, transaction
from django.db.models import Count, Sum
from django.views import View
from django.http import HttpResponse, JsonResponse, HttpResponseBadRequest, FileResponse, Http404
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django_extensions import logging
from docx import Document
from file_management.forms import UploadRecordForm
from file_management.models import StudentFolder, UploadEvent
from accounts.decorators import authorized_roles, authorized_record_user
from accounts.models import User, UserRole, UserRecord, RoleRequest, Log, Setting, Department, Adviser
from .forms import CheckedRecordCommentForm
from .models import AuthorRole, Classification, PSCEDClassification, ConferenceLevel, BudgetType, \
    CollaborationType, Author, Conference, PublicationLevel, Publication, Budget, Collaboration, CheckedRecord, Upload, \
    RecordUpload, RecordType, ResearchRecord, CheckedUpload, RecordUploadStatus, RecordDownloadRequest, \
    RecordAuthenticationPin, CheckedRecordComment
from django.shortcuts import redirect   
from pyexcel_xls import get_data as xls_get
from pyexcel_xlsx import get_data as xlsx_get
from django.utils.datastructures import MultiValueDictKeyError
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from . import forms
from accounts.forms import LoginForm
from records.auxfunctions import *
from axes.models import AccessAttempt, AccessBase
from axes.utils import reset
from django.conf import settings
import magic
import re
import json
from django.shortcuts import redirect, render
import requests
import base64
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime, timedelta
from .models import Subscription, SubscriptionPlan


from accounts.models import User  # Assuming 'accounts' is the app name for the User model
import logging
# Configure the logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

FILE_LENGTH = 5242880

from django.shortcuts import render
from django.http import JsonResponse
from .models import Record

from django.http import JsonResponse

from django.db import connections
from django.http import JsonResponse


    
class SubscriptionPageView(View):
    def get(self, request):
        # Logic to render subscribe.html
        return render(request, 'ipams/subscribe.html')

class PaymentPortalView(View):
    def get(self, request):
        # Logic to render paymentportal.html
        return render(request, 'ipams/paymentportal.html')
    
    
def update_record_tags(request, record_id):
    record = Record.objects.get(pk=record_id)
    ip_is_changed = False
    commercialization_is_changed = False
    community_ext_is_changed = False
    if request.POST.get('ip', 'false') == 'true':
        if not record.is_ip:
            ip_is_changed = True
        record.is_ip = True
    else:
        if record.is_ip:
            ip_is_changed = True
        record.is_ip = False
    if request.POST.get('commercialization', 'false') == 'true':
        if not record.for_commercialization:
            commercialization_is_changed = True
        record.for_commercialization = True
    else:
        if record.for_commercialization:
            commercialization_is_changed = True
        record.for_commercialization = False

    if request.POST.get('community_ext', 'false') == 'true':
        if not record.community_extension:
            community_ext_is_changed = True
        record.community_extension = True
    else:
        if record.community_extension:
            community_ext_is_changed = True
        record.community_extension = False
    record.save()
    status = 'disabled'
    if ip_is_changed:
        if record.is_ip:
            status = 'enabled'
        Log(user=request.user,
            action=f'ip_tag status changed to \"{status}\", record ID: <a href="/dashboard/logs/record/{record_id}">#{record_id}</a>').save()
    if commercialization_is_changed:
        if record.for_commercialization:
            status = 'enabled'
        Log(user=request.user,
            action=f'commercialization_tag status changed to \"{status}\", record ID: <a href="/dashboard/logs/record/{record_id}">#{record_id}</a>').save()

    if community_ext_is_changed:
        if record.community_extension:
            status = 'enabled'
        Log(user=request.user,
            action=f'community_extension_tag status changed to \"{status}\", record ID: <a href="/dashboard/logs/record/{record_id}">#{record_id}</a>').save()
    return {'success': True, 'is-ip': record.is_ip, 'for-commercialization': record.for_commercialization,
            'community-ext': record.community_extension}

import logging

# Function to create a payment link
def create_payment_link_view(request):
    tier = request.GET.get('tier', '')
    price_mapping = {'premium': 10000, 'advanced': 14900, 'free': 10000}  # Corrected prices

    try:
        url = 'https://api.paymongo.com/v1/links'
        payload = {
            'data': {
                'attributes': {
                    'amount': price_mapping.get(tier),  # Convert amount to cents
                    'description': f'Payment for {tier} tier',  # Use tier in description
                    'remarks': 'pay',
                    'status': 'unpaid'  # Add status attribute
                }
            }
        }
        encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
        headers = {
            'Authorization': f'Basic {encoded_api_key}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()  # Raise an error for non-2xx status codes

        data = response.json()
        checkout_url = data['data']['attributes']['checkout_url']
        global stored_link_id
        stored_link_id = data['data']['id']

        return redirect(checkout_url)

    except Exception as error:
        logger.error('Error creating payment link: %s', error)
        return redirect('/')  # Redirect to homepage or error page

@csrf_exempt
def generate_pin_and_save_data_view(request):
    if request.method == 'POST':
        pin = request.POST.get('pin')
        record_id = request.POST.get('record_id')
        user_id = request.POST.get('user_id')
        try:
            # Fetch the 'record' object based on the provided 'record_id'
            record = Record.objects.get(pk=record_id)

            # Get the role_id of the current user
            role_id = request.user.role_id  # Adjust this based on your user model

            # Ensure that 'record' is not None
            if record:
                # Validate the entered PIN against the generated PIN
                generated_pin = pin  # For now, we assume that the PIN sent matches the one generated on the client side

                if pin == generated_pin:
                    # Create a new RecordAuthenticationPin object with the provided 'pin' and 'record' as the foreign key
                    record_auth_pin = RecordAuthenticationPin.objects.create(
                        record=record,
                        pin=pin,
                        user_id=user_id,
                        role_id=role_id,  # Include role_id in the creation of the object
                        email=request.user.email,
                    )

                    # Send the PIN to the user via email
                    send_mail(
                        'Generated PIN',
                        f'Hi {request.user.username}! Your generated PIN is: {pin}',
                        'glajera320@gmail.com',
                        [request.user.email],
                        fail_silently=False,
                    )

                    # Successfully created the RecordAuthenticationPin and sent the email
                    return JsonResponse({'success': True})
                else:
                    # Incorrect PIN entered by the user
                    return JsonResponse({'success': False, 'error_message': 'Incorrect PIN'})
            else:
                # Handle the case where 'record' is not found
                return JsonResponse({'success': False, 'error_message': 'Invalid record'})
        except ObjectDoesNotExist:
            # Handle the case where 'record' is not found
            return JsonResponse({'success': False, 'error_message': 'Record not found'})
        except Exception as e:
            # Log the exception for debugging
            print(f"Error creating RecordAuthenticationPin: {e}")
            # Handle other exceptions here if needed

            return JsonResponse({'success': False, 'error_message': 'Error creating RecordAuthenticationPin'})

    return JsonResponse({'success': False, 'error_message': 'Invalid request method'})


class AuthenticationPinView(View):
    name = 'records/authenticaterecords.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_record_user())
    def get(self, request, record_id):
        checked_records = CheckedRecord.objects.filter(record=Record.objects.get(pk=record_id))
        is_removable = False
        record = Record.objects.get(pk=record_id)
        record_uploads = record.recordupload_set.all()
        for checked_record in checked_records:
            if checked_record.status == 'declined':
                is_removable = True
        self.context['record'] = Record.objects.get(pk=record_id)
        self.context['is_removable'] = is_removable
        self.context['recorduploads'] = record_uploads
        return render(request, self.name, self.context)


# custom function to check the request type since Httpis_ajax(request=request) method is deprecated.
def is_ajax(request):
    return request.META.get('HTTP_X_REQUESTED_WITH') == 'XMLHttpRequest'


class Home(View):
    name = 'records/index.html'

    def get(self, request):
        login_required = request.GET.get('next', False)
        user_roles = UserRole.objects.all()
        departments = Department.objects.all()

        logs = request.session.get('logs', '')
        year_from = request.session.get('year_from', '')
        year_to = request.session.get('year_to', '')

        context = {
            'login_required': login_required,
            'record_form': forms.RecordForm(),
            'login_form': LoginForm(),
            'site_key': settings.GOOGLE_RECAPTCHA_SITE_KEY,
            'test_form': settings.TEST_FORM,
            'user_roles': user_roles,
            'logs': logs,
            'year_from': year_from,
            'year_to': year_to,
            'landing_page': Setting.objects.get(name='landing_page'),

            'departments': departments,
        }
        if logs != '':
            del request.session['logs']
            request.session.modified = True
        return render(request, self.name, context)

    def post(self, request):
        if is_ajax(request=request):
            data = []
            ip_tag = ''
            commercialization_tag = ''
            community_tag = ''
            checked_records = CheckedRecord.objects.filter(status='approved', checked_by__in=Subquery(
                User.objects.filter(role=5).values('pk'))).select_related('record')
            records = Record.objects.filter(pk__in=Subquery(checked_records.values('record_id')))

            # removing accounts
            if request.POST.get('remove-accounts'):
                accounts = request.POST.getlist('accounts[]')
                success = False
                for account_id in accounts:
                    del_account = User.objects.get(pk=int(account_id))
                    if not del_account.is_superuser:
                        del_account.delete()
                        success = True
                return JsonResponse({'success': success})
            # filtering records
            elif request.POST.get('is_filtered') == 'true':
                year_from_filter = request.POST.get('year_from', '0')
                year_to_filter = request.POST.get('year_to', '0')
                classification_filter = request.POST.get('classification')
                psced_classification_filter = request.POST.get('psced_classification')
                author_filter = request.POST.get('author')
                conference_filter = request.POST.get('conference')
                publication_filter = request.POST.get('publication')
                budget_min_filter = request.POST.get('budget_min')
                budget_max_filter = request.POST.get('budget_max')
                collaborator_filter = request.POST.get('collaborator')
                department_filter = request.POST['department']
                ip_filter = request.POST.get('ip_cb')
                commercialization_filter = request.POST.get('commercialization_cb')
                community_filter = request.POST.get('community_cb')

                if (year_from_filter != '' and year_to_filter != ''):
                    records = records.filter(year_accomplished__gte=year_from_filter).filter(
                        year_completed__lte=year_to_filter)
                elif year_from_filter != '':
                    records = records.filter(year_accomplished=year_from_filter)
                elif year_to_filter != '':
                    records = records.filter(year_completed=year_to_filter)

                if ip_filter != '' and commercialization_filter == '' and community_filter == '':
                    records = records.filter(is_ip=True)
                if ip_filter != '' and commercialization_filter != '' and community_filter == '':
                    records = records.filter(Q(is_ip=True) & Q(for_commercialization=True))
                if ip_filter != '' and commercialization_filter == '' and community_filter != '':
                    records = records.filter(Q(is_ip=True) & Q(community_extension=True))
                if commercialization_filter != '' and ip_filter == '' and community_filter == '':
                    records = records.filter(for_commercialization=True)
                if commercialization_filter != '' and ip_filter == '' and community_filter != '':
                    records = records.filter(Q(for_commercialization=True) & Q(community_extension=True))
                if community_filter != '' and commercialization_filter == '' and ip_filter == '':
                    records = records.filter(community_extension=True)
                if ip_filter != '' and commercialization_filter != '' and community_filter != '':
                    records = records.filter(
                        Q(is_ip=True) & Q(for_commercialization=True) & Q(community_extension=True))

                if classification_filter != '':
                    records = records.filter(classification=classification_filter)
                if psced_classification_filter != '':
                    records = records.filter(psced_classification=psced_classification_filter)
                if department_filter != '':
                    records = records.filter(pk__in=Subquery(UserRecord.objects.filter(Q(user__in=Subquery(
                        Student.objects.filter(course__in=Subquery(Course.objects.filter(department__in=Subquery(
                            Department.objects.filter(name=department_filter).values('pk'))).values('pk'))).values(
                            'pk'))) | Q(user__in=Subquery(Adviser.objects.filter(
                        department__in=Subquery(Department.objects.filter(name=department_filter).values('pk'))).values(
                        'pk')))).values('record')))
                if author_filter != '':
                    records = records.filter(
                        pk__in=Author.objects.filter(name__contains=author_filter).values('record_id'))
                if conference_filter != '':
                    records = records.filter(
                        pk__in=Conference.objects.filter(title__contains=conference_filter).values('record_id'))
                if publication_filter != '':
                    publications = Publication.objects.filter(name=publication_filter)
                    if publications.exists():
                        records = records.filter(publication=publications.first())
                if budget_min_filter != "" or budget_max_filter != "":
                    min = 0
                    if budget_min_filter != "":
                        min = float(budget_min_filter)
                    if budget_max_filter != "":
                        max = float(budget_max_filter)
                        records = records.filter(
                            pk__in=Budget.objects.values('record_id').annotate(Sum('budget_allocation')).filter(
                                budget_allocation__sum__range=(min, max)).values('record_id'))
                    else:
                        records = records.filter(
                            pk__in=Budget.objects.values('record_id').annotate(Sum('budget_allocation')).filter(
                                Q(budget_allocation__sum__gte=min)).values('record_id'))
                if collaborator_filter != '':
                    records = records.filter(
                        Q(pk__in=Collaboration.objects.filter(industry__contains=collaborator_filter).values(
                            'record_id')) | Q(
                            pk__in=Collaboration.objects.filter(institution__contains=collaborator_filter).values(
                                'record_id')))
            # accounts role change
            elif request.POST.get('role-change') == 'true':
                accounts = request.POST.getlist('accounts[]')
                accounts_str = ''

                role_id = int(request.POST.get('role-radio'))
                for account_id in accounts:
                    user = User.objects.get(pk=int(account_id))
                    user.role = UserRole.objects.get(pk=role_id)
                    user.save()
                    if account_id == accounts[0]:
                        accounts_str += user.username
                    else:
                        accounts_str += f', {user.username}'
                    RoleRequest.objects.filter(user=user).delete()
                Log(user=request.user,
                    action=f'accounts: {accounts_str} account_role changed to \"{user.role}\" by: {request.user.username}').save()
                # roleRequestNotify(request.user.id, user.id)
                roleRequestApproved(request, request.user.id, user.id)
            # setting datatable records
            for record in records:
                record_conference = Conference.objects.filter(record=record.id)
                if record_conference.exists():
                    for i in record_conference:
                        record_conference_title = i.title
                        record_conference_venue = i.venue
                        if i.conference_level is not None:
                            record_conference_level = i.conference_level.name
                        else:
                            record_conference_level = ''
                else:
                    record_conference_level = ''
                    record_conference_title = ''
                    record_conference_venue = ''

                record_publication = Publication.objects.filter(record=record.id)
                if record_publication.exists():
                    for i in record_publication:
                        record_publication_name = i.name
                        record_publication_isbn = i.isbn
                        record_publication_issn = i.issn
                        record_publication_isi = i.isi
                        if i.publication_level is not None:
                            record_publication_level = i.publication_level.name
                        else:
                            record_publication_level = ''
                else:
                    record_publication_name = ''
                    record_publication_isbn = ''
                    record_publication_issn = ''
                    record_publication_isi = ''
                    record_publication_level = ''

                record_budget = Budget.objects.filter(record=record.id)
                if record_budget.exists():
                    for i in record_budget:
                        record_budget_allocation = i.budget_allocation
                        record_funding_source = i.funding_source
                        if i.budget_type is not None:
                            record_budget_type = i.budget_type.name
                        else:
                            record_budget_type = ''
                else:
                    record_budget_type = ''
                    record_budget_allocation = ''
                    record_funding_source = ''

                record_collaboration = Collaboration.objects.filter(record=record.id)
                if record_collaboration.exists():
                    for i in record_collaboration:
                        record_collaboration_institution = i.institution
                        record_collaboration_industry = i.industry
                        if i.collaboration_type is not None:
                            record_collaboration_type = i.collaboration_type.name
                        else:
                            record_collaboration_type = ''
                else:
                    record_collaboration_type = ''
                    record_collaboration_institution = ''
                    record_collaboration_industry = ''

                if record.is_ip == True:
                    ip_tag = 'Intellectual Property'
                if record.for_commercialization == True:
                    commercialization_tag = 'Commercialization'
                if record.community_extension == True:
                    community_tag = 'Community Extension'

                if request.user.is_anonymous:
                    data.append([
                        record.pk,
                        # additional
                        record.representative,
                        record_conference_level,
                        record_conference_title,
                        record_conference_venue,
                        record_publication_name,
                        record_publication_isbn,
                        record_publication_issn,
                        record_publication_isi,
                        record_publication_level,
                        record_budget_type,
                        record_budget_allocation,
                        record_funding_source,
                        record_collaboration_type,
                        record_collaboration_institution,
                        record_collaboration_industry,
                        ip_tag,
                        commercialization_tag,
                        community_tag,
                        record.abstract,
                        # visible
                        '<a class="record-link" href="/record/' + str(
                            record.pk) + '">' + record.title + '</a>',
                        record.year_accomplished,
                        record.classification.name,
                        record.psced_classification.name,
                    ])
                else:
                    if request.user.role.pk > 2:
                        data.append([
                            record.pk,
                            # additional
                            record.representative,
                            record_conference_level,
                            record_conference_title,
                            record_conference_venue,
                            record_publication_name,
                            record_publication_isbn,
                            record_publication_issn,
                            record_publication_isi,
                            record_publication_level,
                            record_budget_type,
                            record_budget_allocation,
                            record_funding_source,
                            record_collaboration_type,
                            record_collaboration_institution,
                            record_collaboration_industry,
                            ip_tag,
                            commercialization_tag,
                            community_tag,
                            record.abstract,
                            # visible
                            '<a href="/record/' + str(
                                record.pk) + '">' + record.title + '</a>',
                            record.year_accomplished,
                            record.year_completed,
                            record.classification.name,
                            record.psced_classification.name,
                        ])
                    else:
                        data.append([
                            record.pk,
                            # additional
                            record.representative,
                            record_conference_level,
                            record_conference_title,
                            record_conference_venue,
                            record_publication_name,
                            record_publication_isbn,
                            record_publication_issn,
                            record_publication_isi,
                            record_publication_level,
                            record_budget_type,
                            record_budget_allocation,
                            record_funding_source,
                            record_collaboration_type,
                            record_collaboration_institution,
                            record_collaboration_industry,
                            ip_tag,
                            commercialization_tag,
                            community_tag,
                            record.abstract,
                            # visible
                            '<a href="/record/' + str(
                                record.pk) + '">' + record.title + '</a>',
                            record.year_accomplished,
                            record.classification.name,
                            record.psced_classification.name,
                        ])

            return JsonResponse({"data": data})


class PublishedRecords(View):
    name = 'records/published.html'

    def get(self, request):
        login_required = request.GET.get('next', False)
        user_roles = UserRole.objects.all()
        departments = Department.objects.all()

        logs = request.session.get('logs', '')
        year_from = request.session.get('year_from', '')
        year_to = request.session.get('year_to', '')
        context = {
            'login_required': login_required,
            'record_form': forms.RecordForm(),
            'login_form': LoginForm(),
            'site_key': settings.GOOGLE_RECAPTCHA_SITE_KEY,
            'test_form': settings.TEST_FORM,
            'user_roles': user_roles,
            'logs': logs,
            'year_from': year_from,
            'year_to': year_to,
            'landing_page': Setting.objects.get(name='landing_page'),

            'departments': departments,
        }
        if logs != '':
            del request.session['logs']
            request.session.modified = True
        return render(request, self.name, context)

    def post(self, request):
        if is_ajax(request=request):
            data = []
            ip_tag = ''
            commercialization_tag = ''
            community_tag = ''
            checked_records = CheckedRecord.objects.filter(status='approved', checked_by__in=Subquery(
                User.objects.filter(role=5).values('pk'))).select_related('record')
            records = Record.objects.filter(pk__in=Subquery(checked_records.values('record_id')))

            # filtering records
            if request.POST.get('is_filtered') == 'true':
                year_from_filter = request.POST.get('year_from', '0')
                year_to_filter = request.POST.get('year_to', '0')
                classification_filter = request.POST.get('classification')
                psced_classification_filter = request.POST.get('psced_classification')
                author_filter = request.POST.get('author')
                conference_filter = request.POST.get('conference')
                publication_filter = request.POST.get('publication')
                budget_min_filter = request.POST.get('budget_min')
                budget_max_filter = request.POST.get('budget_max')
                collaborator_filter = request.POST.get('collaborator')
                department_filter = request.POST['department']
                ip_filter = request.POST.get('ip_cb')
                commercialization_filter = request.POST.get('commercialization_cb')
                community_filter = request.POST.get('community_cb')

                if (year_from_filter != '' and year_to_filter != ''):
                    records = records.filter(year_accomplished__gte=year_from_filter).filter(
                        year_completed__lte=year_to_filter)
                elif year_from_filter != '':
                    records = records.filter(year_accomplished=year_from_filter)
                elif year_to_filter != '':
                    records = records.filter(year_completed=year_to_filter)

                if ip_filter != '' and commercialization_filter == '' and community_filter == '':
                    records = records.filter(is_ip=True)
                if ip_filter != '' and commercialization_filter != '' and community_filter == '':
                    records = records.filter(Q(is_ip=True) & Q(for_commercialization=True))
                if ip_filter != '' and commercialization_filter == '' and community_filter != '':
                    records = records.filter(Q(is_ip=True) & Q(community_extension=True))
                if commercialization_filter != '' and ip_filter == '' and community_filter == '':
                    records = records.filter(for_commercialization=True)
                if commercialization_filter != '' and ip_filter == '' and community_filter != '':
                    records = records.filter(Q(for_commercialization=True) & Q(community_extension=True))
                if community_filter != '' and commercialization_filter == '' and ip_filter == '':
                    records = records.filter(community_extension=True)
                if ip_filter != '' and commercialization_filter != '' and community_filter != '':
                    records = records.filter(
                        Q(is_ip=True) & Q(for_commercialization=True) & Q(community_extension=True))

                if classification_filter != '':
                    records = records.filter(classification=classification_filter)
                if psced_classification_filter != '':
                    records = records.filter(psced_classification=psced_classification_filter)
                if department_filter != '':
                    records = records.filter(pk__in=Subquery(UserRecord.objects.filter(Q(user__in=Subquery(
                        Student.objects.filter(course__in=Subquery(Course.objects.filter(department__in=Subquery(
                            Department.objects.filter(name=department_filter).values('pk'))).values('pk'))).values(
                            'pk'))) | Q(user__in=Subquery(Adviser.objects.filter(
                        department__in=Subquery(Department.objects.filter(name=department_filter).values('pk'))).values(
                        'pk')))).values('record')))
                if author_filter != '':
                    records = records.filter(
                        pk__in=Author.objects.filter(name__contains=author_filter).values('record_id'))
                if conference_filter != '':
                    records = records.filter(
                        pk__in=Conference.objects.filter(title__contains=conference_filter).values('record_id'))
                if publication_filter != '':
                    publications = Publication.objects.filter(name=publication_filter)
                    if publications.exists():
                        records = records.filter(publication=publications.first())
                if budget_min_filter != "" or budget_max_filter != "":
                    min = 0
                    if budget_min_filter != "":
                        min = float(budget_min_filter)
                    if budget_max_filter != "":
                        max = float(budget_max_filter)
                        records = records.filter(
                            pk__in=Budget.objects.values('record_id').annotate(Sum('budget_allocation')).filter(
                                budget_allocation__sum__range=(min, max)).values('record_id'))
                    else:
                        records = records.filter(
                            pk__in=Budget.objects.values('record_id').annotate(Sum('budget_allocation')).filter(
                                Q(budget_allocation__sum__gte=min)).values('record_id'))
                if collaborator_filter != '':
                    records = records.filter(
                        Q(pk__in=Collaboration.objects.filter(industry__contains=collaborator_filter).values(
                            'record_id')) | Q(
                            pk__in=Collaboration.objects.filter(institution__contains=collaborator_filter).values(
                                'record_id')))

            # setting datatable records
            for record in records:
                record_conference = Conference.objects.filter(record=record.id)
                if record_conference.exists():
                    for i in record_conference:
                        record_conference_title = i.title
                        record_conference_venue = i.venue
                        if i.conference_level is not None:
                            record_conference_level = i.conference_level.name
                        else:
                            record_conference_level = ''
                else:
                    record_conference_level = ''
                    record_conference_title = ''
                    record_conference_venue = ''

                record_publication = Publication.objects.filter(record=record.id)
                if record_publication.exists():
                    for i in record_publication:
                        record_publication_name = i.name
                        record_publication_isbn = i.isbn
                        record_publication_issn = i.issn
                        record_publication_isi = i.isi
                        if i.publication_level is not None:
                            record_publication_level = i.publication_level.name
                        else:
                            record_publication_level = ''
                else:
                    record_publication_name = ''
                    record_publication_isbn = ''
                    record_publication_issn = ''
                    record_publication_isi = ''
                    record_publication_level = ''

                record_budget = Budget.objects.filter(record=record.id)
                if record_budget.exists():
                    for i in record_budget:
                        record_budget_allocation = i.budget_allocation
                        record_funding_source = i.funding_source
                        if i.budget_type is not None:
                            record_budget_type = i.budget_type.name
                        else:
                            record_budget_type = ''
                else:
                    record_budget_type = ''
                    record_budget_allocation = ''
                    record_funding_source = ''

                record_collaboration = Collaboration.objects.filter(record=record.id)
                if record_collaboration.exists():
                    for i in record_collaboration:
                        record_collaboration_institution = i.institution
                        record_collaboration_industry = i.industry
                        if i.collaboration_type is not None:
                            record_collaboration_type = i.collaboration_type.name
                        else:
                            record_collaboration_type = ''
                else:
                    record_collaboration_type = ''
                    record_collaboration_institution = ''
                    record_collaboration_industry = ''

                if record.is_ip == True:
                    ip_tag = 'Intellectual Property'
                if record.for_commercialization == True:
                    commercialization_tag = 'Commercialization'
                if record.community_extension == True:
                    community_tag = 'Community Extension'

                if request.user.is_anonymous:
                    data.append([
                        record.pk,
                        # additional
                        record.representative,
                        record_conference_level,
                        record_conference_title,
                        record_conference_venue,
                        record_publication_name,
                        record_publication_isbn,
                        record_publication_issn,
                        record_publication_isi,
                        record_publication_level,
                        record_budget_type,
                        record_budget_allocation,
                        record_funding_source,
                        record_collaboration_type,
                        record_collaboration_institution,
                        record_collaboration_industry,
                        ip_tag,
                        commercialization_tag,
                        community_tag,
                        record.abstract,
                        # visible
                        '<a href="/record/' + str(
                            record.pk) + '">' + record.title + '</a>',
                        record.year_accomplished,
                        record.classification.name,
                        record.psced_classification.name,
                    ])
                else:
                    if request.user.role.pk > 2:
                        data.append([
                            record.pk,
                            # additional
                            record.representative,
                            record_conference_level,
                            record_conference_title,
                            record_conference_venue,
                            record_publication_name,
                            record_publication_isbn,
                            record_publication_issn,
                            record_publication_isi,
                            record_publication_level,
                            record_budget_type,
                            record_budget_allocation,
                            record_funding_source,
                            record_collaboration_type,
                            record_collaboration_institution,
                            record_collaboration_industry,
                            ip_tag,
                            commercialization_tag,
                            community_tag,
                            record.abstract,
                            # visible
                            '<a href="/record/' + str(
                                record.pk) + '">' + record.title + '</a>',
                            record.year_accomplished,
                            record.year_completed,
                            record.classification.name,
                            record.psced_classification.name,
                        ])
                    else:
                        data.append([
                            record.pk,
                            # additional
                            record.representative,
                            record_conference_level,
                            record_conference_title,
                            record_conference_venue,
                            record_publication_name,
                            record_publication_isbn,
                            record_publication_issn,
                            record_publication_isi,
                            record_publication_level,
                            record_budget_type,
                            record_budget_allocation,
                            record_funding_source,
                            record_collaboration_type,
                            record_collaboration_institution,
                            record_collaboration_industry,
                            ip_tag,
                            commercialization_tag,
                            community_tag,
                            record.abstract,
                            # visible
                            '<a href="/record/' + str(
                                record.pk) + '">' + record.title + '</a>',
                            record.year_accomplished,
                            record.classification.name,
                            record.psced_classification.name,
                        ])

            return JsonResponse({"data": data})

        # Manage documents table


class ManageAccounts(View):
    name = 'records/manage_accounts.html'

    def get(self, request):
        login_required = request.GET.get('next', False)
        user_roles = UserRole.objects.all()
        departments = Department.objects.all()

        logs = request.session.get('logs', '')
        year_from = request.session.get('year_from', '')
        year_to = request.session.get('year_to', '')
        context = {
            'login_required': login_required,
            'record_form': forms.RecordForm(),
            'login_form': LoginForm(),
            'site_key': settings.GOOGLE_RECAPTCHA_SITE_KEY,
            'test_form': settings.TEST_FORM,
            'user_roles': user_roles,
            'logs': logs,
            'year_from': year_from,
            'year_to': year_to,
            'landing_page': Setting.objects.get(name='landing_page'),

            'departments': departments,
        }
        if logs != '':
            del request.session['logs']
            request.session.modified = True
        return render(request, self.name, context)

    def post(self, request):
        # removing accounts
        if request.POST.get('remove-accounts'):
            accounts = request.POST.getlist('accounts[]')
            success = False
            for account_id in accounts:
                del_account = User.objects.get(pk=int(account_id))
                if not del_account.is_superuser:
                    del_account.delete()
                    success = True
            return JsonResponse({'success': success})

        # accounts role change
        elif request.POST.get('role-change') == 'true':
            accounts = request.POST.getlist('accounts[]')
            accounts_str = ''

            role_id = int(request.POST.get('role-radio'))
            for account_id in accounts:
                user = User.objects.get(pk=int(account_id))
                user.role = UserRole.objects.get(pk=role_id)
                user.save()
                if account_id == accounts[0]:
                    accounts_str += user.username
                else:
                    accounts_str += f', {user.username}'
                RoleRequest.objects.filter(user=user).delete()
            Log(user=request.user,
                action=f'accounts: {accounts_str} account_role changed to \"{user.role}\" by: {request.user.username}').save()
            # roleRequestNotify(request.user.id, user.id)
            roleRequestApproved(request, request.user.id, user.id)


class ViewManageDocuments(View):
    name = 'records/dashboard/manage_documents.html'
    record_uploads = RecordUpload.objects.all()
    record_upload_status = RecordUploadStatus.objects.all()
    context = {
        'record_uploads': record_uploads,
        'record_upload_status': record_upload_status,
    }

    @method_decorator(authorized_roles(roles=['adviser', 'ktto', 'rdco', 'tbi', 'itso']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        return render(request, self.name, self.context)

    def post(self, request):
        if is_ajax(request=request):
            if request.POST.get('status_change', False) == 'true':
                record_upload = RecordUpload.objects.get(pk=int(request.POST.get('record_upload_id', 0)))
                record_upload.record_upload_status = RecordUploadStatus.objects.get(pk=request.POST.get('status', 0))
                record_upload.save()
                Log(user=request.user,
                    action=f'{record_upload.upload.name}_document status changed to \"{record_upload.record_upload_status}\", record ID: <a href="/dashboard/logs/record/{record_upload.record.pk}">#{record_upload.record.pk}</a>').save()
                return JsonResponse({'success': True})
            else:
                data = []
                record_uploads = RecordUpload.objects.all()
                if request.POST.get('is-filter', '0') == '1' and request.POST.get('record-upload-status', '0') != '0':
                    record_uploads = record_uploads.filter(record_upload_status=RecordUploadStatus.objects.get(
                        pk=request.POST.get('record-upload-status', '0')))
                for record_upload in record_uploads:
                    record = record_upload.record
                    approval = "Not Yet Approved"
                    if record_upload.is_approved:
                        approval = "Approved"
                    data.append([record_upload.pk,
                                 record_upload.upload.name,
                                 f'<a href="/download/document/{record_upload.pk}">{record_upload.filename}</a>',
                                 f'<a href="/dashboard/manage/records/{record.pk}">{record.title}</a>',
                                 f'{record_upload.record.record_type.name}',
                                 f'<a href="/file_management/personal_files/folders/records/attachments/{record_upload.record.pk}" target="_blank">View All</a>',
                                 f'{record_upload.date_uploaded.strftime("%Y-%m-%d %H:%M:%S")}',
                                 approval,
                                 record_upload.record_upload_status.name,
                                 f'<button type="button" class="btn base-btn base-bg-secondary" onclick="onStatusChangeClick({record_upload.pk}, {record_upload.record_upload_status.pk});">Change</button>',
                                 ])

        return JsonResponse({"data": data})


# Manage documents template
class ViewManageDocumentsRecord(View):
    name = 'records/dashboard/manage_documents_record.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
        'is_owner': True,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_record_user())
    def get(self, request, record_id):
        owners = UserRecord.objects.filter(record=Record.objects.get(pk=record_id))
        self.context['owners'] = owners
        checked_records = CheckedRecord.objects.filter(record=Record.objects.get(pk=record_id))
        adviser_checked = {'status': 'pending'}
        ktto_checked = {'status': 'pending'}
        rdco_checked = {'status': 'pending'}
        role_checked = False
        record = Record.objects.get(pk=record_id)
        research_record = ResearchRecord.objects.filter(Q(proposal=record) | Q(research=record)).first()
        for checked_record in checked_records:
            if checked_record.checked_by.role.id == 3:
                adviser_checked = {'status': checked_record.status, 'content': checked_record}
            if checked_record.checked_by.role.id == 4 or checked_record.checked_by.role.id == 7:
                ktto_checked = {'status': checked_record.status, 'content': checked_record}
            if checked_record.checked_by.role.id == 5:
                rdco_checked = {'status': checked_record.status, 'content': checked_record}
            if checked_record.checked_by.role.id == request.user.role.pk:
                role_checked = True
        self.context['adviser_checked'] = adviser_checked
        self.context['ktto_checked'] = ktto_checked
        self.context['rdco_checked'] = rdco_checked
        self.context['role_checked'] = role_checked
        self.context['record'] = record
        self.context['is_removable'] = True
        self.context['research_record'] = research_record
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # removing record
            if request.POST.get('remove', 'false') == 'true':
                del_record = Record.objects.get(pk=record_id)
                del_record.abstract_file.delete()
                del_record_uploads = RecordUpload.objects.filter(record=del_record)
                for del_record_upload in del_record_uploads:
                    del_record_upload.file.delete()
                del_record.delete()
                return JsonResponse({'success': True})
            # updating record tags
            elif request.POST.get('tags_update', 'false') == 'true':
                return JsonResponse(update_record_tags(request, record_id))
            # get uploaded document data
            elif request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'doc-status': record_upload.record_upload_status.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            # POSTING COMMENTS
            elif request.POST.get('post_comment', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                comment = request.POST.get('comment', '')
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                CheckedUpload(comment=comment, checked_by=request.user, record_upload=record_upload).save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False})


class ViewRecordFile(View):
    name = 'records/view_record_file.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'tbi', 'itso']))
    def get(self, request, record_id):

        if request.user.role.pk == 2:
            userRecord = UserRecord.objects.filter(record=record_id).first()
            if request.user.pk != userRecord.user.pk:
                context = {'permission': 'restrict'}
                return render(request, self.name, context)

        try:
            record = Record.objects.get(pk=record_id)
            record_path = record.abstract_file.path
            record_file = record.abstract_file
            filename = record.abstract_filename
            mime = magic.Magic(mime=True)
            content_type = mime.from_file(record_path)
            if content_type == 'application/pdf':
                response = FileResponse(record_file, content_type='application/pdf')
                return response
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                record_url = record.abstract_file.url  # Get the file URL
                context = {'record_url': record_url, 'permission': 'allow'}
                return render(request, self.name, context)
        except FileNotFoundError:
            raise Http404()


class ViewRecord(View):
    name = 'records/view.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_record_user())
    def get(self, request, record_id):
        view_record_controller = ViewRecordController()
        data = view_record_controller.getRecordInfo(request, record_id)
        self.context.update(data)
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # updating record tags
            if request.POST.get('tags_update', 'false') == 'true':
                return JsonResponse(update_record_tags(request, record_id))

            # send download request
            if request.POST.get('sendRequest'):
                record_id = request.POST.get('recordId')
                user_id = request.POST.get('userId')
                download_request = RecordDownloadRequest(sent_by=User.objects.get(pk=user_id),
                                                         record=Record.objects.get(pk=record_id))
                download_request.save()
                sendDownloadRequest(request, request.user.id, record_id)
                return JsonResponse({'success': True})


            # get uploaded document data
            elif request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'doc-status': record_upload.record_upload_status.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            # POSTING COMMENTS
            elif request.POST.get('post_comment', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                comment = request.POST.get('comment', '')
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                CheckedUpload(comment=comment, checked_by=request.user,
                              record_upload=record_upload).save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False})


class ViewRecordController(View):
    def getRecordInfo(self, request, record_id):
        data = {}

        try:
            record = Record.objects.get(pk=record_id)
        except Record.DoesNotExist:
            return HttpResponseBadRequest({'error': 'Record not found'}, status=404)
        owners = UserRecord.objects.filter(record=record)
        owners_data = [owner.user.username for owner in owners]
        adviser_checked = {'status': 'Pending'}
        ktto_checked = {'status': 'Pending'}
        rdco_checked = {'status': 'Pending'}
        research_record = ResearchRecord.objects.filter(Q(proposal=record) | Q(research=record)).first()
        checked_record = CheckedRecord.objects.filter(record=record)
        for checked_record in checked_record:
            if checked_record.checked_by.role.id == 3:
                adviser_checked = {'status': checked_record.status.capitalize()}    
            if checked_record.checked_by.role.id == 4 or checked_record.checked_by.role.id == 7:
                ktto_checked = {'status': checked_record.status.capitalize()}
            if checked_record.checked_by.role.id == 5:
                rdco_checked = {'status': checked_record.status.capitalize()}
        download_requests = RecordDownloadRequest.objects.filter(record=record_id,sent_by=request.user.id)
        download_request_sent = download_requests.exists()
        download_approval = any(request.is_marked for request in download_requests)
        data['download_request_sent'] = download_request_sent
        data['download_approval'] = download_approval
        data['owners'] = owners_data
        data['recorduploads'] = record.recordupload_set.all()
        data['adviser_status'] = adviser_checked
        data['ktto_status'] = ktto_checked
        data['rdco_status'] = rdco_checked
        data['record'] = record
        data['research_record'] = research_record
        data['is_owner'] = UserRecord.objects.filter(user=request.user, record=record).exists()
        data['is_removable'] = request.user.role.pk > 3 or any(
            cr.status == 'declined' for cr in CheckedRecord.objects.filter(record=record))

        if Conference.objects.filter(record=record_id).exists():
            data['conference'] = 'true'
        else:
            data['conference'] = 'false'
        if Publication.objects.filter(record=record_id).exists():
            data['publication'] = 'true'
        else:
            data['publication'] = 'false'
        if Budget.objects.filter(record=record_id).exists():
            data['budget'] = 'true'
        else:
            data['budget'] = 'false'
        if Collaboration.objects.filter(record=record_id).exists():
            data['collaboration'] = 'true'
        else:
            data['collaboration'] = 'false'

        return data


class UploadController(View):
    def get(self, request):
        def getUploadDocument(record_id, upload_id):
            print(upload_id)
            upload = Upload.objects.get(pk=upload_id)
            record = Record.objects.get(pk=record_id)
            record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
            checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
            comments = []
            checked_bys = []
            checked_dates = []
            for checked_upload in checked_uploads:
                comments.append(checked_upload.comment)
                checked_bys.append(checked_upload.checked_by.username)
                checked_dates.append(checked_upload.date_checked)
            if record_upload is None:
                return JsonResponse({'success': False, 'doc-title': upload.name})
            else:
                return JsonResponse({'success': True,
                                     'doc-title': record_upload.upload.name,
                                     'doc-filename': record_upload.filename,
                                     'is-ip': record_upload.is_ip,
                                     'for-commercialization': record_upload.for_commercialization,
                                     'comments': comments,
                                     'checked_bys': checked_bys,
                                     'checked_dates': checked_dates,
                                     'record-upload-id': record_upload.pk})

        if is_ajax(request):
            if request.path == '/record/uploads/getUploadDocument/':
                record_id = request.GET.get('record_id')
                upload_id = request.GET.get('upload_id')
                return getUploadDocument(record_id, upload_id)
        else:
            return HttpResponseBadRequest("Invalid request")

    def post(self, request):
        def createComment(record_id, upload_id):
            upload = Upload.objects.get(pk=upload_id)
            record = Record.objects.get(pk=record_id)
            comment = request.POST.get('comment', '')
            record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
            CheckedUpload(comment=comment, checked_by=request.user,
                          record_upload=record_upload).save()

            recordComment(request, request.user.id, record.id, UserRecord.objects.get(record=record.id).user.id)
            return JsonResponse({'success': True})

        if is_ajax(request):
            record_id = request.POST.get('record_id')
            upload_id = request.POST.get('upload_id')
            if request.path == '/record/uploads/createComment/':
                return createComment(record_id, upload_id)
            else:
                return HttpResponseBadRequest("Invalid request")
        else:
            return HttpResponseBadRequest("Invalid request")


class EvaluationView(View):
    name = 'records/profile/evaluation.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    checked_record_form = CheckedRecordCommentForm()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
        'checked_record_form': checked_record_form,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['adviser', 'ktto', 'rdco', 'tbi', 'itso']))
    def get(self, request, record_id):
        view_record_controller = ViewRecordController()
        data = view_record_controller.getRecordInfo(request, record_id)
        self.context.update(data)
        return render(request, self.name, self.context)


class EvaluationController(View):

    def post(self, request):
        def removeRecord(record_id):
            reason = request.POST.get('reason')
            record = Record.objects.get(pk=record_id)
            record.is_marked = True
            record.reason = reason
            record.save()
            deleteRecord(request, request.user.id, record.pk, reason)
            return JsonResponse({'success': True})

        def evaluateRecord(record_id):
            message = []
            data = {'status': 'error', 'message': 'There was an error processing your request.'}
            status = request.POST.get('status')
            comment_save_status = False
            if CheckedRecord.objects.filter(checked_by = request.user.id,record = record_id).exists():
                checked_record = CheckedRecord.objects.get(checked_by = request.user.id,record = record_id)
            else:
                checked_record = None

            if request.user.role.id == 3: 
                checked_record_comment_form = CheckedRecordCommentForm(request.POST)
                if checked_record is not None:
                    checked_record = CheckedRecord.objects.get(checked_by = request.user.id,record = record_id)
                    if checked_record_comment_form.is_valid():
                        checked_record_comment = checked_record_comment_form.save(commit=False)
                        checked_record_comment.checked_record = checked_record
                        checked_record_comment.save()
                        checked_record.status = request.POST.get('status')
                        checked_record.save()
                        comment_save_status = True
                        
                    else:
                        message = 'Comment must be given for evaluation to be valid'
                        messages.error(request, message)
                else:
                    if checked_record_comment_form.is_valid():
                        checked_record = CheckedRecord.objects.create(record = Record.objects.get(pk=record_id), status = request.POST.get('status'), checked_by = request.user)
                        checked_record_comment = checked_record_comment_form.save(commit=False)
                        checked_record_comment.checked_record = checked_record
                        checked_record_comment.save()
                        comment_save_status = True
                    else:
                        message = 'Comment must be given for evaluation to be valid'
                        messages.error(request, message)
            
            elif request.user.role.id > 3:
                checked_record_comment_form = CheckedRecordCommentForm(request.POST)
                if checked_record is not None:                
                    if checked_record.checked_by is not request.user:
                        checked_record.checked_by = request.user
                        checked_record.status = status
                        checked_record.save()
                        comment_save_status = True


                        if checked_record_comment_form.is_valid():
                            checked_record_comment = checked_record_comment_form.save(commit=False)
                            checked_record_comment.checked_record = checked_record
                            checked_record_comment.save()
                            
                    else:
                        message = 'Comment must be given for evaluation to be valid'
                        messages.error(request, message)

                else:
                    if checked_record_comment_form.is_valid():
                        checked_record = CheckedRecord.objects.create(record = Record.objects.get(pk=record_id), status = request.POST.get('status'), checked_by = request.user)
                        checked_record_comment = checked_record_comment_form.save(commit=False)
                        checked_record_comment.checked_record = checked_record
                        checked_record_comment.save()
                        comment_save_status = True
                    else:
                        message = 'Comment must be given for evaluation to be valid'
                        messages.error(request, message)

            else:
                message = 'You do not have the necessary role to evaluate records'

            checked_record_status = request.POST.get('status')
            if comment_save_status:
                saveStatus = recordStatus(request, request.user.id, record_id,
                    UserRecord.objects.get(record=record_id).user.id,
                    checked_record_status)
                if saveStatus:
                    if checked_record_status == 'approved':
                        updateTagStatus = update_record_tags(request, record_id)

                        if updateTagStatus:
                            message = 'Your evaluation has been successfully recorded'
                            status = 'success'
                            messages.success(request, message)
                            data = {'status': status, 'message': message,
                                    'checked_record_status': checked_record_status}

                            # Inside the block where the folder is created
                            if request.user.role.id == 5:
                                # Get the current user's ID and role ID
                                student_role_id = 2
                                added_by_user = Record.objects.get(pk=record_id).added_by
                                record = Record.objects.get(pk=record_id)

                                # Create a name for the folder based on user ID (you may need to modify this logic)
                                folder_count = StudentFolder.objects.filter(role__id=student_role_id,
                                                                            user=added_by_user).count()
                                folder_name = f"Project {folder_count + 1}"

                                # Create a StudentFolder with the role set to the student role instance
                                student_role = UserRole.objects.get(pk=student_role_id)
                                student_folder = StudentFolder(
                                    name=folder_name,
                                    user=added_by_user,
                                    record=record,
                                    role=student_role
                                )
                                student_folder.save()

                        else:
                            message = 'There was an error in handling your record Tag'
                            status = 'error'
                            messages.success(request, message)
                    else:
                        message = 'Your evaluation has been successfully recorded'
                        status = 'success'
                        messages.success(request, message)
                        data = {'status': status, 'message': message,
                            'checked_record_status': checked_record_status}
                else:
                    message = 'Your evaluation has been succesfully recorded'
                    status = 'success'
                    messages.success(request, message)
                    data = {'status': status, 'message': message,
                                'checked_record_status': checked_record_status}
            else:
                message = 'There was an error in handling your evaluation, please try again'
                status = 'error'
                messages.error(request, message)
                data = {'status': status, 'message': message,
                            'checked_record_status': checked_record_status}
                
            data['message'] = message
            data['status'] = status
            data['checked_record_status'] = request.POST.get('status')
            return JsonResponse(data)

        if is_ajax(request):
            record_id = request.POST.get('record_id')
            if request.path == '/record/evaluation/removeRecord/':
                return removeRecord(record_id)
            elif request.path == '/record/evaluation/evaluateRecord/':
                return evaluateRecord(record_id)
            else:
                return HttpResponseBadRequest("Invalid request")
        else:
            return HttpResponseBadRequest("Invalid request")


# display all delete record requests to table
def get_all_delete_requests(request):
    if request.method == 'POST':
        data = []
        delete_pendings = Record.objects.filter(is_marked=True)
        for record in delete_pendings:
            data.append([
                record.pk,
                f'<a href="/record/pending/delete/request/{record.pk}">{record.title}</a>',
                record.reason,
            ])

        return JsonResponse({'data': data})


# template view when selecting from pending delete requests table

class PendingDeleteRecordsView(View):
    name = 'records/profile/pending_delete_records.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
    }

    @method_decorator(login_required(login_url='/'))
    def get(self, request, record_id):
        owners = UserRecord.objects.filter(record=Record.objects.get(pk=record_id))
        self.context['owners'] = owners
        checked_records = CheckedRecord.objects.filter(record=Record.objects.get(pk=record_id))
        role_checked = False
        is_owner = False
        is_removable = False
        if request.user.role.pk > 3:
            is_removable = True
        for checked_record in checked_records:
            if checked_record.checked_by.role.id == request.user.role.pk:
                role_checked = True
            if checked_record.status == 'declined':
                is_removable = True
        if UserRecord.objects.filter(user=request.user, record=Record.objects.get(pk=record_id)):
            is_owner = True
        self.context['role_checked'] = role_checked
        self.context['record'] = Record.objects.get(pk=record_id)
        self.context['is_owner'] = is_owner
        self.context['is_removable'] = is_removable
        self.context['reason'] = Record.objects.get(pk=record_id).reason
        self.context['delete_request'] = Record.objects.filter(is_marked=True)
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if request.method == 'POST':
            # removing record
            if 'approvedRequestBtn' in request.POST:
                record_id = request.POST.get('record_number')
                del_record = Record.objects.get(pk=record_id)
                user_record = UserRecord.objects.get(record=Record.objects.get(pk=record_id))
                del_record.abstract_file.delete()
                del_record_uploads = RecordUpload.objects.filter(record=del_record)
                for del_record_upload in del_record_uploads:
                    del_record_upload.file.delete()
                approvedDeleteRecord(request, request.user.id, record_id, user_record.user.pk, del_record.reason)
                del_record.delete()
                return redirect("records-pending")
            else:
                messages.error(request, 'Error encountered while approving record')
                return redirect("records-pending")

        if is_ajax(request=request):
            # get uploaded document data
            if request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            else:
                return JsonResponse({'success': False})

            # return redirect('records-pending')


# template view when selecting from my records table
class MyRecordView(View):
    name = 'records/profile/view_myrecords.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
        'is_owner': True,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_record_user())
    def get(self, request, record_id):
        view_record_controller = ViewRecordController()
        data = view_record_controller.getRecordInfo(request, record_id)
        self.context.update(data)
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # removing record
            if request.POST.get('deleteRecord', 'false') == 'true':
                reason = request.POST.get('reason')
                recordTitle = request.POST.get('recordTitle')
                record = Record.objects.get(title=recordTitle)
                record.is_marked = True
                record.reason = reason
                record.save()
                deleteRecord(request, request.user.id, record.pk, reason)
                return JsonResponse({'success': True})

            # if request.POST.get('remove', 'false') == 'true':
            #     del_record = Record.objects.get(pk=record_id)
            #     del_record.abstract_file.delete()
            #     del_record_uploads = RecordUpload.objects.filter(record=del_record)
            #     for del_record_upload in del_record_uploads:
            #         del_record_upload.file.delete()
            #     del_record.delete()
            #     return JsonResponse({'success': True})

            # resubmitting
            if request.POST.get('resubmit', 'false') == 'true':
                record=Record.objects.get(pk=record_id)
                CheckedRecord.objects.filter(record=record).update(status='pending')
                recipient = record.adviser.id
                resubmission(request, request.user.id, record_id, recipient)
                return JsonResponse({'success': True})
            # updating record tags
            elif request.POST.get('tags_update', 'false') == 'true':
                return JsonResponse(update_record_tags(request, record_id))
            # get uploaded document data
            elif request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'doc-status': record_upload.record_upload_status.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            # POSTING COMMENTS
            elif request.POST.get('post_comment', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                comment = request.POST.get('comment', '')
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                CheckedUpload(comment=comment, checked_by=request.user, record_upload=record_upload).save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False})


class ApprovedRecordView(View):
    name = 'records/profile/view_approved.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
    }

    @method_decorator(login_required(login_url='/'))
    def get(self, request, record_id):
        view_record_controller = ViewRecordController()
        data = view_record_controller.getRecordInfo(request, record_id)
        self.context.update(data)
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # removing record
            if request.POST.get('remove', 'false') == 'true':
                del_record = Record.objects.get(pk=record_id)
                del_record.abstract_file.delete()
                del_record.delete()
                return JsonResponse({'success': True})
            # updating record tags
            elif request.POST.get('tags_update', 'false') == 'true':
                return JsonResponse(update_record_tags(request, record_id))
            # get uploaded document data
            elif request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            # POSTING COMMENTS
            elif request.POST.get('post_comment', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                comment = request.POST.get('comment', '')
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                CheckedUpload(comment=comment, checked_by=request.user,
                              record_upload=record_upload).save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False})


class DeclinedRecordView(View):
    name = 'records/profile/view_declined.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
    }

    @method_decorator(login_required(login_url='/'))
    def get(self, request, record_id):
        view_record_controller = ViewRecordController()
        data = view_record_controller.getRecordInfo(request, record_id)
        self.context.update(data)
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # removing record
            if request.POST.get('remove', 'false') == 'true':
                del_record = Record.objects.get(pk=record_id)
                del_record.abstract_file.delete()
                del_record.delete()
                return JsonResponse({'success': True})
            # get uploaded document data
            elif request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            # POSTING COMMENTS
            elif request.POST.get('post_comment', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                comment = request.POST.get('comment', '')
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                CheckedUpload(comment=comment, checked_by=request.user,
                              record_upload=record_upload).save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False})


class AddRecordView(View):
    name = 'records/add.html'

    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'itso', 'tbi']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        author_roles = AuthorRole.objects.all()
        record_types = RecordType.objects.all()
        upload = Upload.objects.all()

        context = {
            'author_roles': author_roles,
            'record_types': record_types,
            'record_form': forms.RecordForm(),
            'publication_form': forms.PublicationForm(),
            'author_roles': AuthorRole.objects.all(),
            'conference_levels': ConferenceLevel.objects.all(),
            'budget_types': BudgetType.objects.all(),
            'collaboration_types': CollaborationType.objects.all(),
            'uploads': upload,
        }
        return render(request, self.name, context)

from datetime import datetime, timedelta

class AddRecordController(View):
    user = None
    uploadStatus = None
    notificationStatus = None
    ownerList = []
    adviserList = []

    def post(self, request):
        def getUser():
            user = User.objects.get(pk=request.user.pk)
            return user

        def createNotification(user, adviser, record):
            try:
                user = User.objects.get(pk=user)
                adviser = User.objects.get(pk=adviser)
                record = Record.objects.get(pk=record)

                if record.record_type_id == 1 or record.record_type_id == 2:
                    notif_type = NotificationType.objects.get(pk=3)
                elif record.record_type_id == 3:
                    notif_type = NotificationType.objects.get(pk=4)

                if user.role.name == 'Student':
                    course = Student.objects.get(user=user.id).course.name
                    role = UserRole.objects.get(pk=2)
                elif user.role.name == 'Adviser':
                    course = ''
                    role = UserRole.objects.get(pk=3)
                elif user.role.name == 'KTTO':
                    course = ''
                    role = UserRole.objects.get(pk=4)
                elif user.role.name == 'RDCO':
                    course = ''
                    role = UserRole.objects.get(pk=5)
                elif user.role.name == 'TBI':
                    course = ''
                    role = UserRole.objects.get(pk=7)

                notification = Notification(user=user, course=course, role=role, recipient=adviser, record=record,
                                            notif_type=notif_type,
                                            is_read=False, date_created=datetime.now())

                notification.save()
                return True
            except Exception as e:
                print(str(e))
                return False

        def getUserList():
            if request.POST.get("get_user_tags", 'false') == 'true':
                users = User.objects.values('username', 'pk')
                data = {"users": list(users)}
                return JsonResponse(data)
            else:
                return HttpResponseBadRequest("Invalid request")

        def getAdviserList():
            if request.POST.get("get_user_tags", 'false') == 'true':
                advisers = User.objects.filter(role__in=[3, 4, 5]).values('username', 'pk')
                data = {"advisers": list(advisers)}
                return JsonResponse(data)
            else:
                return HttpResponseBadRequest("Invalid request")

        def createRecord():
            user = getUser()
            try:

                record = Record()

                # set record requirements
                record.setTitle(request.POST.get('title'))
                record.setYearCompleted(request.POST.get('year_completed'))
                record.setYearAccomplished(request.POST.get('year_accomplished'))
                abstract_content = request.POST.get('abstract_content')
                cleaned_abstract = re.sub(r'<\/?p>', '', abstract_content)  # Remove <p> and </p> tags
                record.setAbstract(cleaned_abstract)
                record.setClassification(Classification.objects.get(pk=request.POST.get('classification')))
                record.setPscedClassification(
                    PSCEDClassification.objects.get(pk=request.POST.get('psced_classification')))
                record.setRecordType(RecordType.objects.get(pk=request.POST.get('record_type')))
                # set file requirements
                record.setAbstractFile(request.FILES.get('abstract_file'))
                record.setAbstractFileSize(record.abstract_file.size)
                record.setAbstractFileName(record.abstract_file.name)

                # set owners and evaluators
                owners = json.loads(request.POST.get('owners-id'))
                adviser = json.loads(request.POST.get('adviser-id'))
                record.setRepresentative(f'{request.user.first_name} {request.user.last_name}')
                record.setAdviser(User.objects.get(pk=adviser[0]['id']))

                # Set the added_by field to the currently logged-in user (assuming the user is a student)
                if user.role.pk == 2:
                    record.added_by = user

                # If User role is student add record code
                if user.role.pk == 2:
                    student = Student.objects.select_related('course__department__college').get(user=user)
                    year = str(datetime.now().year)[2:]
                    serial = record.pk
                    college = student.course.department.college.code
                    department = student.course.department.code
                    record.representative = f'{request.user.first_name} {request.user.last_name}'
                    record.code = f'{year}-{serial}-{college}-{department}-{request.user.last_name.upper()}'

                # Save the record
                record.save()

                recommendations = request.POST.get('recommendations', 'No recommendations provided')
                
                # Mapping for psced_classification.pk to corresponding names
                classification_mapping = {
                    14: 'Education Science and Teacher Training',
                    18: 'Fine and Applied Arts',
                    22: 'Humanities',
                    26: 'Religion and Theology',
                    30: 'Social and Behavioral Sciences',
                    34: 'Business Administration and Related',
                    38: 'Law and Jurisprudence',
                    42: 'Natural Science',
                    46: 'Mathematics',
                    47: 'IT Related Disciplines',
                    50: 'Medical and Allied',
                    52: 'Trade, Craft and Industrial',
                    54: 'Engineering and Technology',
                    58: 'Architecture and Town Planning',
                    62: 'Agriculture, Forestry and Fisheries',
                    66: 'Home Economics',
                    78: 'Service Trades',
                    84: 'Mass Communication and Documentation',
                    89: 'Other Disciplines'
                }

                # Get the classification name from the mapping
                classification_name = classification_mapping.get(record.psced_classification.pk, None)

                if classification_name is None:
                    raise ValueError("Invalid PSCED classification ID")
                # Execute the insert query using the 'nalc' database connection
                with connections['nalc'].cursor() as cursor:
                    cursor.execute("""
                        INSERT INTO backend_researchpaper 
                        (title, abstract, year, classification, author, recommendations)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    """, [
                        record.title,
                        cleaned_abstract,
                        record.year_accomplished,
                        classification_name,
                        record.representative,
                        recommendations,
                    ])


                # Save under research record
                if record.record_type.pk == 1:
                    ResearchRecord(proposal=record).save()

                # Save record
                for owner in owners:
                    UserRecord(user=User.objects.get(pk=int(owner['id'])), record=record).save()

                # Upload records
                for upload in Upload.objects.all():
                    if record.record_type.pk == upload.record_type.pk:
                        if request.FILES.get(f'upload-{upload.pk}', None):
                            file = request.FILES.get(f'upload-{upload.pk}', None)
                            RecordUpload(file=file, filename=file.name, record=record,
                                         upload=upload,
                                         record_upload_status=RecordUploadStatus.objects.get(pk=1)).save()

                author_names = request.POST.getlist('author_names[]', None)
                author_roles = request.POST.getlist('author_roles[]', None)
                for i, author_name in enumerate(author_names):
                    Author(name=author_name, author_role=AuthorRole.objects.get(pk=author_roles[i]),
                           record=record).save()
                if record.record_type.pk == 3:
                    publication_form = forms.PublicationForm(request.POST)
                    if publication_form.is_valid():
                        publication = publication_form.save(commit=False)
                        if publication.name or publication.isbn or publication.isi or publication.issn or publication.publication_level or publication.year_published:
                            publication.record = record
                            publication.save()

                    conference_levels = request.POST.getlist('conference_levels[]', None)
                    conference_titles = request.POST.getlist('conference_titles[]', None)
                    conference_dates = request.POST.getlist('conference_dates[]', None)
                    conference_venues = request.POST.getlist('conference_venues[]', None)
                    budget_types = request.POST.getlist('budget_types[]', None)
                    budget_allocations = request.POST.getlist('budget_allocations[]', None)
                    funding_sources = request.POST.getlist('funding_sources[]', None)
                    industries = request.POST.getlist('industries[]', None)
                    institutions = request.POST.getlist('institutions[]', None)
                    collaboration_types = request.POST.getlist('collaboration_types[]', None)

                    for i, conference_title in enumerate(conference_titles):
                        Conference(title=conference_title,
                                   conference_level=ConferenceLevel.objects.get(pk=conference_levels[i]),
                                   date=conference_dates[i], venue=conference_venues[i], record=record).save()

                    for i, budget_type in enumerate(budget_types):
                        Budget(budget_type=BudgetType.objects.get(pk=budget_types[i]),
                               budget_allocation=budget_allocations[i],
                               funding_source=funding_sources[i], record=record).save()
                    for i, collaboration_type in enumerate(collaboration_types):
                        Collaboration(collaboration_type=CollaborationType.objects.get(pk=collaboration_types[i]),
                                      industry=industries[i], institution=institutions[i], record=record).save()

                notify = createNotification(user.id, record.adviser.id, record.id)
                if notify:
                    messege = 'Record has been submitted and your adviser has been notified'
                    status = 'success'
                else:
                    messege = 'Record has been submitted but there was an error in notifying your adviser, please contact your adviser'
                    status = 'notifError'

                messages.success(request, messege)
                data = {'status': status, 'messege': messege}

                return JsonResponse(data)

            except Exception as e:
                messege = 'There was an error in uploading record, Error:' + str(e)
                status = 'error'
                messages.error(request, messege)
                data = {'status': status, 'messege': messege}

                return JsonResponse(data)

        if is_ajax(request):
            if request.path == '/add/getUserList/':
                return getUserList()
            elif request.path == '/add/getAdviserList/':
                return getAdviserList()
            elif request.path == '/add/createRecord/':
                return createRecord()
            else:
                return HttpResponseBadRequest("Invalid request")
        else:
            logging.debug("Received an invalid request")
            return redirect("/")


class AddResearch(View):
    name = 'records/add_research.html'
    author_roles = AuthorRole.objects.all()
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    record_types = RecordType.objects.all()
    record_form = forms.RecordForm()
    publication_form = forms.PublicationForm()
    uploads = Upload.objects.all()

    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'itso', 'tbi']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request, research_record_id):
        proposal_record = ResearchRecord.objects.get(pk=research_record_id).proposal
        context = {
            'author_roles': self.author_roles,
            'conference_levels': self.conference_levels,
            'budget_types': self.budget_types,
            'collaboration_types': self.collaboration_types,
            'record_types': self.record_types,
            'record_form': self.record_form,
            'publication_form': self.publication_form,
            'uploads': self.uploads,
            'proposal_record': proposal_record,
            'research_record_id': research_record_id,
        }
        return render(request, self.name, context)

    def post(self, request, research_record_id):
        error_messages = []
        record_form = forms.RecordForm(request.POST, request.FILES)
        proposal_record = ResearchRecord.objects.get(pk=research_record_id).proposal
        if is_ajax(request=request):
            if request.POST.get("get_user_tags", 'false') == 'true':
                users = []
                advisers = []
                for user in User.objects.all():
                    users.append({'value': user.username, 'id': user.pk})
                for user in User.objects.filter(role__in=[3, 4, 5]):
                    advisers.append({'value': user.username, 'id': user.pk})
                return JsonResponse({'users': users, 'advisers': advisers})
        if record_form.is_valid() and not is_ajax(request=request):
            record = record_form.save(commit=False)
            file_is_valid = True
            file = record_form.cleaned_data.get('abstract_file', False)
            # check uploaded file size if valid
            if file and file.size > FILE_LENGTH:
                file_is_valid = False
            # saving record to database
            else:
                owners = json.loads(request.POST.get('owners-id'))
                adviser = json.loads(request.POST.get('adviser-id'))
                record.adviser = User.objects.get(pk=adviser[0]['id'])
                record.record_type = RecordType.objects.get(pk=2)

                if file is not None:
                    record.abstract_filesize = record.abstract_file.size
                    record.abstract_filename = record.abstract_file.name
                    # upload_blob(settings.GS_BUCKET_NAME, record.abstract_file, record.abstract_file.name)
                    # record.abstract_file = None

                record.save()
                research_record = ResearchRecord.objects.get(pk=research_record_id)
                research_record.research = record
                research_record.save()
                # patent search files check
                for upload in Upload.objects.all():
                    if request.FILES.get(f'upload-{upload.pk}', None):
                        file = request.FILES.get(f'upload-{upload.pk}', None)
                        # upload_blob(settings.GS_BUCKET_NAME, file, file.name)
                        # RecordUpload(file=None, filename=file.name, record=record,
                        #                              upload=upload).save()
                        RecordUpload(file=file, filename=file.name, record=record, upload=upload).save()
                for owner in owners:
                    UserRecord(user=User.objects.get(pk=int(owner['id'])), record=record).save()
            if record is not None and file_is_valid:
                publication_form = forms.PublicationForm(request.POST)
                if publication_form.is_valid():
                    publication = publication_form.save(commit=False)
                    publication.record = record
                    publication.save()
                author_names = request.POST.getlist('author_names[]', None)
                author_roles = request.POST.getlist('author_roles[]', None)
                conference_levels = request.POST.getlist('conference_levels[]', None)
                conference_titles = request.POST.getlist('conference_titles[]', None)
                conference_dates = request.POST.getlist('conference_dates[]', None)
                conference_venues = request.POST.getlist('conference_venues[]', None)

                budget_types = request.POST.getlist('budget_types[]', None)
                budget_allocations = request.POST.getlist('budget_allocations[]', None)
                funding_sources = request.POST.getlist('funding_sources[]', None)
                industries = request.POST.getlist('industries[]', None)
                institutions = request.POST.getlist('institutions[]', None)
                collaboration_types = request.POST.getlist('collaboration_types[]', None)
                for i, author_name in enumerate(author_names):
                    Author(name=author_name, author_role=AuthorRole.objects.get(pk=author_roles[i]),
                           record=record).save()

                for i, conference_title in enumerate(conference_titles):
                    Conference(title=conference_title,
                               conference_level=ConferenceLevel.objects.get(pk=conference_levels[i]),
                               date=conference_dates[i], venue=conference_venues[i], record=record).save()

                for i, budget_type in enumerate(budget_types):
                    Budget(budget_type=BudgetType.objects.get(pk=budget_types[i]),
                           budget_allocation=budget_allocations[i],
                           funding_source=funding_sources[i], record=record).save()
                for i, collaboration_type in enumerate(collaboration_types):
                    Collaboration(collaboration_type=CollaborationType.objects.get(pk=collaboration_types[i]),
                                  industry=industries[i], institution=institutions[i], record=record).save()
                newRecordAdded(request, request.user.id, record.adviser.id, record.id)
                return redirect('records-index')
            elif not file_is_valid:
                error = {'title': 'Unable to save record',
                         'body': 'The file cannot be more than 5 MB'}
                error_messages.append(error)
            else:
                error = {'title': 'Unable to save record',
                         'body': 'A record with the same record information already exists'}
                error_messages.append(error)
        else:
            error_messages.append({'title': 'Unable to save record',
                                   'body': 'Some fields contains invalid values while trying to save the record'})
        context = {
            'author_roles': self.author_roles,
            'conference_levels': self.conference_levels,
            'budget_types': self.budget_types,
            'collaboration_types': self.collaboration_types,
            'record_form': record_form,
            'publication_form': self.publication_form,
            'proposal_record': proposal_record,
            'research_record_id': research_record_id,
            'error_messages': error_messages,
        }
        return render(request, self.name, context)


class AddResearchController(View):
    def post(self, request):
        def getUser():
            user = User.objects.get(pk=request.user.pk)
            return user

        def createNotification(user, adviser, record):
            try:
                user = User.objects.get(pk=user)
                adviser = User.objects.get(pk=adviser)
                record = Record.objects.get(pk=record)

                if record.record_type_id == 1 or record.record_type_id == 2:
                    notif_type = NotificationType.objects.get(pk=3)
                elif record.record_type_id == 3:
                    notif_type = NotificationType.objects.get(pk=4)

                if user.role.name == 'Student':
                    course = Student.objects.get(user=user.id).course.name
                    role = UserRole.objects.get(pk=2)
                elif user.role.name == 'Adviser':
                    course = ''
                    role = UserRole.objects.get(pk=3)
                elif user.role.name == 'KTTO':
                    course = ''
                    role = UserRole.objects.get(pk=4)
                elif user.role.name == 'RDCO':
                    course = ''
                    role = UserRole.objects.get(pk=5)
                elif user.role.name == 'TBI':
                    course = ''
                    role = UserRole.objects.get(pk=7)

                notification = Notification(user=user, course=course, role=role, recipient=adviser, record=record,
                                            notif_type=notif_type,
                                            is_read=False, date_created=dt.now())

                notification.save()
                return True
            except Exception as e:
                print(str(e))
                return False

        def createThesis():
            user = getUser()
            try:

                record = Record()
                research_record_id = request.POST.get('researchId')
                proposal_record = ResearchRecord.objects.get(pk=research_record_id).proposal

                # set record requirements
                record.setTitle(request.POST.get('title'))
                record.setYearCompleted(request.POST.get('year_completed'))
                record.setYearAccomplished(request.POST.get('year_accomplished'))
                record.setAbstract(request.POST.get('abstract_content'))
                record.setClassification(Classification.objects.get(pk=request.POST.get('classification')))
                record.setPscedClassification(
                    PSCEDClassification.objects.get(pk=request.POST.get('psced_classification')))
                record.setRecordType(RecordType.objects.get(pk=2))
                # set file requirements
                record.setAbstractFile(request.FILES.get('abstract_file'))
                record.setAbstractFileSize(record.abstract_file.size)
                record.setAbstractFileName(record.abstract_file.name)

                # set owners and evaluators
                owners = json.loads(request.POST.get('owners-id'))
                adviser = json.loads(request.POST.get('adviser-id'))
                record.setRepresentative(f'{request.user.first_name} {request.user.last_name}')
                record.setAdviser(User.objects.get(pk=adviser[0]['id']))

                # If User role is student add record code
                if user.role.pk == 2:
                    student = Student.objects.select_related('course__department__college').get(user=user)
                    year = str(datetime.datetime.now().year)[2:]
                    serial = record.pk
                    college = student.course.department.college.code
                    department = student.course.department.code
                    record.representative = f'{request.user.first_name} {request.user.last_name}'
                    record.code = f'{year}-{serial}-{college}-{department}-{request.user.last_name.upper()}'

                # Save the record
                record.save()

                research_record = ResearchRecord.objects.get(pk=research_record_id)
                research_record.research = record
                research_record.save()

                # Save under research record
                if record.record_type.pk == 1:
                    ResearchRecord(proposal=record).save()

                # Save record
                for owner in owners:
                    UserRecord(user=User.objects.get(pk=int(owner['id'])), record=record).save()

                # Upload records
                for upload in Upload.objects.all():
                    if record.record_type.pk == upload.record_type.pk:
                        if request.FILES.get(f'upload-{upload.pk}', None):
                            file = request.FILES.get(f'upload-{upload.pk}', None)
                            RecordUpload(file=file, filename=file.name, record=record,
                                         upload=upload,
                                         record_upload_status=RecordUploadStatus.objects.get(pk=1)).save()

                author_names = request.POST.getlist('author_names[]', None)
                author_roles = request.POST.getlist('author_roles[]', None)
                for i, author_name in enumerate(author_names):
                    Author(name=author_name, author_role=AuthorRole.objects.get(pk=author_roles[i]),
                           record=record).save()

                publication_form = forms.PublicationForm(request.POST)
                if publication_form.is_valid():
                    publication = publication_form.save(commit=False)
                    if publication.name or publication.isbn or publication.isi or publication.issn or publication.publication_level or publication.year_published:
                        publication.record = record
                        publication.save()

                conference_levels = request.POST.getlist('conference_levels[]', None)
                conference_titles = request.POST.getlist('conference_titles[]', None)
                conference_dates = request.POST.getlist('conference_dates[]', None)
                conference_venues = request.POST.getlist('conference_venues[]', None)
                budget_types = request.POST.getlist('budget_types[]', None)
                budget_allocations = request.POST.getlist('budget_allocations[]', None)
                funding_sources = request.POST.getlist('funding_sources[]', None)
                industries = request.POST.getlist('industries[]', None)
                institutions = request.POST.getlist('institutions[]', None)
                collaboration_types = request.POST.getlist('collaboration_types[]', None)

                for i, conference_title in enumerate(conference_titles):
                    Conference(title=conference_title,
                               conference_level=ConferenceLevel.objects.get(pk=conference_levels[i]),
                               date=conference_dates[i], venue=conference_venues[i], record=record).save()

                for i, budget_type in enumerate(budget_types):
                    Budget(budget_type=BudgetType.objects.get(pk=budget_types[i]),
                           budget_allocation=budget_allocations[i],
                           funding_source=funding_sources[i], record=record).save()
                for i, collaboration_type in enumerate(collaboration_types):
                    Collaboration(collaboration_type=CollaborationType.objects.get(pk=collaboration_types[i]),
                                  industry=industries[i], institution=institutions[i], record=record).save()

                notify = createNotification(user.id, record.adviser.id, record.id)
                if notify:
                    messege = 'Thesis Record has been submitted and your adviser has been notified'
                    status = 'success'
                else:
                    messege = 'Thesis Record has been submitted but there was an error in notifying your adviser, please contact your adviser'
                    status = 'notifError'

                messages.success(request, messege)
                data = {'status': status, 'messege': messege}

                return JsonResponse(data)

            except Exception as e:
                messege = 'There was an error in uploading Thesis record, Error:' + str(e)
                status = 'error'
                data = {'status': status, 'messege': messege}

                return JsonResponse(data)

        if is_ajax(request):
            if request.path == '/add/createThesis/':
                return createThesis()
            else:
                logging.debug("Received request for invalid request 1")
                return HttpResponseBadRequest("Invalid request")
        else:
            logging.debug("Received an invalid request")


class Edit(View):
    name = 'records/edit.html'
    author_roles = AuthorRole.objects.all()
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    record_types = RecordType.objects.all()
    # record_form = forms.RecordForm()
    record_form = forms.EditRecordForm()
    uploads = Upload.objects.all()

    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'itso', 'tbi']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request, record_id):
        record = Record.objects.get(pk=record_id)
        authors = Author.objects.filter(record=record)
        conferences = Conference.objects.filter(record=record)
        budgets = Budget.objects.filter(record=record)
        collaborations = Collaboration.objects.filter(record=record)
        # record_form = forms.RecordForm(instance=record)
        record_form = forms.EditRecordForm(instance=record)
        try:
            publication_form = forms.PublicationForm(instance=Publication.objects.get(record=record))
            publication_name = Publication.objects.get(record=record).name
        except Publication.DoesNotExist:
            publication_form = forms.PublicationForm()
            publication_name = ""
        record_uploads = record.recordupload_set.all()
        context = {
            'author_roles': self.author_roles,
            'conference_levels': self.conference_levels,
            'budget_types': self.budget_types,
            'collaboration_types': self.collaboration_types,
            'record_types': self.record_types,
            'record_form': record_form,
            'publication_form': publication_form,
            'publication_name': publication_name,
            'record': record,
            'authors': authors,
            'conferences': conferences,
            'budgets': budgets,
            'collaborations': collaborations,
            'uploads': self.uploads,
            'recorduploads': record_uploads,
            'filename': record.abstract_filename,
        }
        return render(request, self.name, context)

    def post(self, request, record_id):
        error_messages = []
        record_instance = Record.objects.get(pk=record_id)
        # record_form = forms.RecordForm(request.POST, request.FILES, instance=Record.objects.get(pk=record_id))
        record_form = forms.EditRecordForm(request.POST, request.FILES, instance=Record.objects.get(pk=record_id))

        if record_form.is_valid():
            record = record_form.save(commit=False)
            if record is None:
                record = record_instance
            record.record_type = record_instance.record_type
            file_is_valid = True
            file = record_form.cleaned_data.get('abstract_file', False)
            # check abstract file size if valid
            if file and file.size > FILE_LENGTH:
                file_is_valid = False

            # saving record to database
            else:
                if file and file != record_instance.abstract_filename:
                    print(file, ' != ', record_instance.abstract_filename)
                    # delete_blob(settings.GS_BUCKET_NAME, record_instance.abstract_filename)
                    record.abstract_filesize = record.abstract_file.size
                    record.abstract_filename = record.abstract_file.name
                    # upload_blob(settings.GS_BUCKET_NAME, record.abstract_file, record.abstract_file.name)
                    # record.abstract_file = None

                record.save()
                # documents search files check
                # for upload in Upload.objects.all():
                #    if request.FILES.get(f'upload-{upload.pk}', None):
                #        record_upload = RecordUpload.objects.filter(record=record, upload=upload).first()
                #        file = request.FILES.get(f'upload-{upload.pk}', None)

                #        if record_upload is not None:
                # if file and file != record_upload.filename:
                #     print(file, ' != ', record_upload.filename)
                #     delete_blob(settings.GS_BUCKET_NAME, record_upload.filename)
                #     upload_blob(settings.GS_BUCKET_NAME, file, file.name)
                #            if record_upload.record_upload_status.pk not in [2, 3, 5]:
                #                record_upload.file = file
                # record_upload.file = None
                #                record_upload.filename = file.name
                #                record_upload.save()
                #            else:
                #                messages.error(request, 'Cannot be updated, document has been processed')
                #        else:
                # upload_blob(settings.GS_BUCKET_NAME, file, file.name)
                # RecordUpload(file=None, filename=file.name, record=record,
                #                          upload=upload, record_upload_status=RecordUploadStatus.objects.get(pk=1)).save()
                #            RecordUpload(file=file, filename=file.name, record=record,
                #                         upload=upload,
                #                         record_upload_status=RecordUploadStatus.objects.get(pk=1)).save()

            if record is not None and file_is_valid:
                try:
                    publication_form = forms.PublicationForm(instance=Publication.objects.get(record=record))
                except Publication.DoesNotExist:
                    publication_form = forms.PublicationForm()
                if publication_form.is_valid():
                    publication = publication_form.save(commit=False)
                    publication.record = record
                    publication.save()
                author_names = request.POST.getlist('author_names[]', None)
                author_roles = request.POST.getlist('author_roles[]', None)
                conference_levels = request.POST.getlist('conference_levels[]', None)
                conference_titles = request.POST.getlist('conference_titles[]', None)
                conference_dates = request.POST.getlist('conference_dates[]', None)
                conference_venues = request.POST.getlist('conference_venues[]', None)

                budget_types = request.POST.getlist('budget_types[]', None)
                budget_allocations = request.POST.getlist('budget_allocations[]', None)
                funding_sources = request.POST.getlist('funding_sources[]', None)
                industries = request.POST.getlist('industries[]', None)
                institutions = request.POST.getlist('institutions[]', None)
                collaboration_types = request.POST.getlist('collaboration_types[]', None)

                Author.objects.filter(record=record).delete()
                for i, author_name in enumerate(author_names):
                    Author(name=author_name, author_role=AuthorRole.objects.get(pk=author_roles[i]),
                           record=record).save()

                Conference.objects.filter(record=record).delete()
                for i, conference_title in enumerate(conference_titles):
                    Conference(title=conference_title,
                               conference_level=ConferenceLevel.objects.get(pk=conference_levels[i]),
                               date=conference_dates[i], venue=conference_venues[i], record=record).save()

                Budget.objects.filter(record=record).delete()
                for i, budget_type in enumerate(budget_types):
                    Budget(budget_type=BudgetType.objects.get(pk=budget_types[i]),
                           budget_allocation=budget_allocations[i],
                           funding_source=funding_sources[i], record=record).save()
                Collaboration.objects.filter(record=record).delete()
                for i, collaboration_type in enumerate(collaboration_types):
                    Collaboration(collaboration_type=CollaborationType.objects.get(pk=collaboration_types[i]),
                                  industry=industries[i], institution=institutions[i], record=record).save()
                newRecordAdded(request, request.user.id, record.adviser.id, record.id)
                return JsonResponse({'success': 1})
            elif not file_is_valid:
                error = {'title': 'Unable to save record',
                         'body': 'The file cannot be more than 5 MB'}
                error_messages.append(error)
            else:
                error = {'title': 'Unable to save record',
                         'body': 'A record with the same record information already exists'}
                error_messages.append(error)
        else:
            error_messages.append({'title': 'Unable to save record',
                                   'body': 'Some fields contains invalid values while trying to save the record'})
        context = {
            'author_roles': self.author_roles,
            'conference_levels': self.conference_levels,
            'budget_types': self.budget_types,
            'collaboration_types': self.collaboration_types,
            'record_form': record_form,
            'record': record_instance,
            'publication_form': self.publication_form,
            'error_messages': error_messages,
        }
        return JsonResponse({'success': 0})


class ParseExcel(View):
    name = 'records/import_record.html'

    def get(self, request):
        return render(request, self.name)

    def post(self, request):
        row_count = 5
        count = 0
        error_count = 0
        now = datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        logs = now + '\n'
        try:
            excel_file = request.FILES['file']
            data = {}
            if str(excel_file).split('.')[-1] == 'xls':
                data = xls_get(excel_file, column_limit=50)
            elif str(excel_file).split('.')[-1] == 'xlsx':
                data = xlsx_get(excel_file, column_limit=50)
            data = data['ResearchProductivity'][6:][0:]
            for d in data:
                row_count += 1
                if d[0] != 'end of records':
                    representative = d[0]
                    title = d[1]
                    year_accomplished = d[2]
                    classification = 1
                    psced_classification = d[5]
                    if not d[3]:
                        classification = 2
                    conference_level = 1
                    conference_title = d[9]
                    conference_date = d[14]
                    conference_venue = d[15]
                    if d[11]:
                        conference_level = 2
                    elif d[12]:
                        conference_level = 3
                    elif d[13]:
                        conference_level = 4
                    record_len = len(Record.objects.filter(title=title, year_accomplished=year_accomplished,
                                                           classification=Classification.objects.get(pk=classification),
                                                           psced_classification=PSCEDClassification.objects.get(
                                                               pk=psced_classification)))
                    if record_len == 0:
                        record = Record(title=title, year_accomplished=year_accomplished,
                                        classification=Classification.objects.get(pk=classification),
                                        psced_classification=PSCEDClassification.objects.get(pk=psced_classification),
                                        record_type=RecordType.objects.get(pk=3), representative=representative)
                        record.save()
                        UserRecord(record=record, user=request.user).save()
                    else:
                        logs += f'\nDuplicate entry on row "{row_count}"'
                        error_count += 1
                        continue
                    Conference(title=conference_title,
                               conference_level=ConferenceLevel.objects.get(pk=conference_level),
                               date=conference_date, venue=conference_venue, record=record).save()
                    if d[16]:
                        publication_name = d[23]
                        sn_list = "".join(d[24].split()).split(',')
                        isbn = ''
                        issn = ''
                        isi = ''
                        for sn in sn_list:
                            if sn.upper().find('ISBN:') >= 0:
                                isbn = sn.replace('ISBN:', '')
                            elif sn.upper().find('ISSN:') >= 0:
                                issn = sn.replace('ISSN:', '')
                            elif sn.upper().find('ISI:') >= 0:
                                isi = sn.replace('ISI:', '')
                        publication_level = 1
                        if d[20]:
                            publication_level = 2
                        elif d[21]:
                            publication_level = 3
                        elif d[22]:
                            publication_level = 4
                        year_published = d[18]
                        Publication(name=publication_name, isbn=isbn, issn=issn, isi=isi,
                                    publication_level=PublicationLevel.objects.get(pk=publication_level),
                                    year_published=year_published, record=record).save()
                    else:
                        Publication(record=record).save()
                    if d[25]:
                        budget_type = 1
                        budget_allocation = d[30]
                        funding_source = d[31]
                        if d[28]:
                            budget_type = 2
                        elif d[20]:
                            budget_type = 3
                        Budget(budget_type=BudgetType.objects.get(pk=budget_type),
                               budget_allocation=budget_allocation,
                               funding_source=funding_source, record=record).save()
                    if d[32]:
                        industry = d[34]
                        institution = d[35]
                        collaboration_type = 1
                        if len(d) >= 38:
                            if d[37]:
                                collaboration_type = 2

                        elif len(d) >= 39:
                            if d[38]:
                                collaboration_type = 3
                        Collaboration(collaboration_type=CollaborationType.objects.get(pk=collaboration_type),
                                      industry=industry, institution=institution, record=record).save()
                    count += 1
                else:
                    break
            messages.success(request, f'{count} records imported!')
            messages.error(request,
                           f'{error_count} records contains errors and cannot be imported. See Upload page for the logs...')
            request.session[
                'logs'] = logs + f'\nTotal records found: {row_count - 6}, Success: {count}, Errors: {error_count}'
        except (MultiValueDictKeyError, KeyError, ValueError, OSError):
            messages.error(request, "Some rows have invalid values")
            print('Multivaluedictkeyerror/KeyError/ValueError/OSError')
        except (DataError, ValidationError):
            messages.error(request, "The form is invalid")
            print('DataError/ValidationError')
        return redirect('records-index')


@authorized_roles(roles=['adviser', 'ktto', 'rdco', 'itso', 'tbi'])
def download_format(request):
    fl_path = '/media'
    filename = 'data.xlsx'
    fl = open('media/data.xlsx', 'rb')
    mime_type = mimetypes.guess_type(fl_path)
    response = HttpResponse(fl, content_type=mime_type)
    response['Content-Disposition'] = "attachment; filename=%s" % filename
    return response


@authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'itso', 'tbi'])
def download_abstract(request, record_id):
    record = Record.objects.get(pk=record_id)
    # filename = record.abstract_file.name.split('/')[-1]
    filename = record.abstract_filename
    response = HttpResponse(record.abstract_file, content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename=%s' % filename
    return response

    # download abstract from google cloud storage
    # record = Record.objects.get(pk=record_id)
    # # source_blob_name = 'abstract/'+record.abstract_filename
    # source_blob_name = record.abstract_filename
    # link = download_blob(settings.GS_BUCKET_NAME, source_blob_name)
    # return HttpResponseRedirect(link)


# documents in upload tab
@authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'itso', 'tbi'])
def download_document(request, record_upload_id):
    record_upload = RecordUpload.objects.get(pk=record_upload_id)
    # filename = record_upload.file.name.split('/')[-1]
    filename = record_upload.filename
    response = HttpResponse(record_upload.file, content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename=%s' % filename
    return response

    # download document from google cloud storage
    # record_upload = RecordUpload.objects.get(pk=record_upload_id)
    # # source_blob_name = 'abstract/'+record.abstract_filename
    # source_blob_name = record_upload.filename
    # link = download_blob(settings.GS_BUCKET_NAME, source_blob_name)
    # return HttpResponseRedirect(link)


# table view of all my records
class MyRecordsView(View):
    template_name = 'records/profile/my_records.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'itso', 'tbi']))
    def get(self, request):
        return render(request, self.template_name)


class MyRecordsController(View):

    def post(self, request):
        def getMyRecordList():
            # Get all records owned by user
            user_records = UserRecord.objects.filter(user=request.user)
            data = []
            # Initialize status badges for records
            for user_record in user_records:
                records = Record.objects.filter(userrecord=user_record)
                for record in records:
                    adviser_checked = f'<div class="badge badge-secondary">pending</div>'
                    ktto_checked = f'<div class="badge badge-secondary">pending</div>'
                    rdco_checked = f'<div class="badge badge-secondary">pending</div>'
                    for checked_record in CheckedRecord.objects.filter(record=record):
                        if checked_record.status == 'approved':
                            badge = 'success'
                        elif checked_record.status == 'declined':
                            badge = 'danger'
                        else:
                            badge = 'secondary'
                        record_status = f'<div class="badge badge-{badge}">{checked_record.status}</div>'
                        if checked_record.checked_by.role.pk == 3:
                            adviser_checked = record_status
                        elif checked_record.checked_by.role.pk == 4 or checked_record.checked_by.role.id == 7:
                            ktto_checked = record_status
                        elif checked_record.checked_by.role.pk == 5:
                            rdco_checked = record_status
                    if record.is_marked == False:
                        data.append([
                            record.pk,
                            '<a href="/record/myrecords/' + str(record.pk) + '">' + record.title + '</a>',
                            '<a href="#" onclick="checkStatusClick(\'adviser\',' + str(
                                record.pk) + ');">' + adviser_checked + '</a>',
                            '<a href="#" onclick="checkStatusClick(\'ktto\',' + str(
                                record.pk) + ');">' + ktto_checked + '</a>',
                            '<a href="#" onclick="checkStatusClick(\'rdco\',' + str(
                                record.pk) + ');">' + rdco_checked + '</a>',
                        ])
                    elif record.is_marked == True:
                        data.append([
                            record.pk,
                            record.title + ' <i class="fa fa-ban" aria-hidden="true" title="Marked for deletion"></i>',
                            '<a href="#" onclick="checkStatusClick(\'adviser\',' + str(
                                record.pk) + ');">' + adviser_checked + '</a>',
                            '<a href="#" onclick="checkStatusClick(\'ktto\',' + str(
                                record.pk) + ');">' + ktto_checked + '</a>',
                            '<a href="#" onclick="checkStatusClick(\'rdco\',' + str(
                                record.pk) + ');">' + rdco_checked + '</a>',
                        ])
            return JsonResponse({"data": data})

        def getComment(role, record_id):
            # get checked record
            if role == 'adviser':
                checked_record = CheckedRecord.objects.filter(record=record_id, checked_by__role__id=3).first()
            elif role == 'ktto':
                checked_record = CheckedRecord.objects.filter(Q(record=record_id), Q(checked_by__role__id=4) | Q(
                    checked_by__role__id=7)).first()
            elif role == 'rdco':
                checked_record = CheckedRecord.objects.filter(record=record_id, checked_by__role__id=5).first()

            if checked_record is None:
                return JsonResponse({}, status=204)

            checked_record_comments = CheckedRecordComment.objects.filter(checked_record = checked_record)
            comments = []
            date_created = []
            for checked_record_comment in checked_record_comments:
                comments.append(checked_record_comment.comment)
                date_created.append(checked_record_comment.date_created)
            data = {"status": checked_record.status,
                    "username": checked_record.checked_by.username,
                    "comment": comments,
                    "date_created": date_created}
            return JsonResponse(data)

        if is_ajax(request):
            if request.path == '/records/user/getMyRecordsList/':
                return getMyRecordList()
            elif request.path == '/records/user/getComment/':
                role = request.POST['role']
                id = request.POST['recordId']
                return getComment(role, id)
            else:
                return HttpResponseBadRequest("Invalid request")
        else:
            return HttpResponseBadRequest("Invalid request")


# Pending records table
class PendingRecordsView(View):
    template_name = 'records/profile/pending_records.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'tbi']))
    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        if request.user.role.id == 3:
            #get all records except those with checked
            #adviser_pending_records = Record.objects.filter(adviser=request.user.pk).exclude(
            #    pk__in=Subquery(adviser_exclude.values('record').distinct())).values('pk', 'title')
            records_with_adviser = Record.objects.filter(adviser=request.user.pk)
            adviser_pending_records = records_with_adviser.filter(
                Q(checkedrecord__isnull=True) | Q(checkedrecord__checked_by=request.user.pk, checkedrecord__status='pending')
            ).distinct().values('pk', 'title')
            tuples_adviser_pending_records = [tuple(a.values()) for a in adviser_pending_records]

            data = []
            for row in tuples_adviser_pending_records:
                data.append([
                    row[0],
                    '<a href="/record/evaluation/' + str(row[0]) + '">' + row[1] + '</a>',
                ])
        elif request.user.role.id == 4 or request.user.role.id == 7:
            ktto_exclude = CheckedRecord.objects.select_related('record').filter(
                Q(checked_by__in=Subquery(User.objects.filter(role=4).values('pk'))) | Q(
                    checked_by__in=Subquery(User.objects.filter(role=7).values('pk')))).exclude(status='pending')
            ktto_include = CheckedRecord.objects.select_related('record').filter(status='approved',
                                                                                 checked_by__in=Subquery(
                                                                                     User.objects.filter(role=3).values(
                                                                                         'pk')))
            ktto_pending_records = Record.objects.filter(pk__in=Subquery(ktto_include.values('record'))).exclude(
                pk__in=Subquery(ktto_exclude.values('record'))).values('pk', 'title', 'is_marked')
            tuples_ktto_pending_records = [tuple(k.values()) for k in ktto_pending_records]

            data = []
            for row in tuples_ktto_pending_records:
                if row[2] == False:
                    data.append([
                        row[0],
                        f'<a href="/record/evaluation/{row[0]}">{row[1]}</a>'
                    ])
                else:
                    data.append([
                        row[0],
                        row[1] + ' <i class="fa fa-ban" aria-hidden="true" title="Marked for deletion"></i>'
                    ])

        elif request.user.role.id == 5:
            rdco_exclude = CheckedRecord.objects.select_related('record').filter(
                checked_by__in=Subquery(User.objects.filter(role=5).values('pk'))).exclude(status='pending')
            rdco_include = CheckedRecord.objects.select_related('record').filter(
                Q(checked_by__in=Subquery(User.objects.filter(role=4).values('pk'))) | Q(
                    checked_by__in=Subquery(User.objects.filter(role=7).values('pk'))), status='approved')
            rdco_pending_records = Record.objects.filter(pk__in=Subquery(rdco_include.values('record'))).exclude(
                pk__in=Subquery(rdco_exclude.values('record'))).values('pk', 'title', 'is_marked')
            tuples_rdco_pending_records = [tuple(r.values()) for r in rdco_pending_records]

            data = []
            for row in tuples_rdco_pending_records:
                if row[2] == False:
                    data.append([
                        row[0],
                        f'<a href="/record/evaluation/{row[0]}">{row[1]}</a>'
                    ])
                else:
                    data.append([
                        row[0],
                        row[1] + ' <i class="fa fa-ban" aria-hidden="true" title="Marked for deletion"></i>'
                    ])
        return JsonResponse({"data": data})


class PendingAttachmentsView(View):
    template_name = 'records/profile/pending_attachments.html'

    record_uploads = RecordUpload.objects.filter(is_approved=False)

    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        # Get only the records where is_approved is 0 or False
        records = CheckedUpload.objects.filter(comment='')

        data = []
        for record in records:
            data.append({
                'id': record.id,
                'file': record.record_upload.file.url if record.record_upload.file else '',
                'filename': record.record_upload.filename,
                'upload': record.record_upload.upload.name,
                'record': record.record_upload.record.title,
                'record_upload_status': str(record.record_upload.record_upload_status),
                'is_ip': record.record_upload.is_ip,
                'for_commercialization': record.record_upload.for_commercialization,
                'date_uploaded': record.record_upload.date_uploaded.strftime('%Y-%m-%d %H:%M:%S'),
                'version': record.record_upload.version,
            })

        return JsonResponse({'data': data})


def pending_attachments_view(request):
    if request.method == 'GET':
        # Retrieve the list of record uploads from the database
        record_uploads = RecordUpload.objects.filter(is_approved=False)

        # Render the template with the record uploads data
        return render(request, 'records/profile/pending_attachments.html', {'record_uploads': record_uploads})
    elif request.method == 'POST':
        # Get only the records where is_approved is 0 or False
        checked_uploads = CheckedUpload.objects.filter(comment='')

        data = []
        for record in checked_uploads:
            if record.record_upload:
                data.append({
                    'id': record.id,
                    'record_upload_id': record.record_upload.id,
                    'file': record.record_upload.file.url if record.record_upload.file else '',
                    'filename': record.record_upload.filename,
                    'upload': record.record_upload.upload.name,
                    'record': record.record_upload.record.title,
                    'record_upload_status': str(record.record_upload.record_upload_status),
                    'is_ip': record.record_upload.is_ip,
                    'for_commercialization': record.record_upload.for_commercialization,
                    'date_uploaded': record.record_upload.date_uploaded.strftime('%Y-%m-%d %H:%M:%S'),
                    'version': record.record_upload.version,
                    'is_approved': record.record_upload.is_approved,
                })

        return JsonResponse({'data': data})


class PendingRequestsView(View):
    template_name = 'records/profile/pending_requests.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['ktto', 'rdco', 'tbi']))
    def get(self, request):
        return render(request, self.template_name)


class ApprovedRecordsView(View):
    template_name = 'records/profile/approved_records.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'tbi']))
    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        checked_records = CheckedRecord.objects.filter(checked_by=request.user, status='approved')
        data = []
        for checked_record in checked_records:
            data.append([
                checked_record.record.pk,
                f'<a href="/record/approved/{checked_record.record.pk}">{checked_record.record.title}</a>'
            ])
        return JsonResponse({'data': data})


class ApprovedAttachmentsView(View):
    template_name = 'records/profile/approved_attachments.html'

    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        # Get all records without filtering based on approval status
        records = RecordUpload.objects.filter(is_approved=1)

        data = []
        for record in records:
            data.append({
                'id': record.id,
                'file': record.file.url if record.file else '',
                'filename': record.filename,
                'upload': record.upload.name,
                'record': record.record.title,
                'record_upload_status': str(record.record_upload_status),
                'is_ip': record.is_ip,
                'for_commercialization': record.for_commercialization,
                'date_uploaded': record.date_uploaded.strftime('%Y-%m-%d %H:%M:%S'),
                'version': record.version,
                'is_approved': record.is_approved,
            })

        return JsonResponse({'data': data})


def approve_attachments_view(request):
    # Retrieve the list of record uploads from the database
    record_uploads = RecordUpload.objects.filter(is_approved=True)

    # Render the template with the record uploads data
    return render(request, 'records/profile/approved_attachments.html', {'record_uploads': record_uploads})


def decline_attachments_view(request):
    # Retrieve the list of record uploads from the database
    record_uploads = RecordUpload.objects.filter(is_declined=True)

    # Render the template with the record uploads data
    return render(request, 'records/profile/declined_attachments.html', {'record_uploads': record_uploads})


class DeclinedAttachmentsView(View):
    template_name = 'records/profile/declined_attachments.html'

    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        # Get all records without filtering based on approval status
        records = RecordUpload.objects.filter(is_declined=True)

        data = []
        for record in records:
            data.append({
                'id': record.id,
                'file': record.file.url if record.file else '',
                'filename': record.filename,
                'upload': record.upload.name,
                'record': record.record.title,
                'record_upload_status': str(record.record_upload_status),
                'is_ip': record.is_ip,
                'for_commercialization': record.for_commercialization,
                'date_uploaded': record.date_uploaded.strftime('%Y-%m-%d %H:%M:%S'),
                'version': record.version,
                'is_declined': record.is_declined,
            })

        return JsonResponse({'data': data})


class DeclinedRecordsView(View):
    template_name = 'records/profile/declined_records.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['student', 'adviser', 'ktto', 'rdco', 'tbi']))
    def get(self, request):
        return render(request, self.template_name)

    def post(self, request):
        checked_records = CheckedRecord.objects.filter(checked_by=request.user, status='declined')
        data = []
        for checked_record in checked_records:
            data.append([
                checked_record.record.pk,
                f'<a href="/record/declined/{checked_record.record.pk}">{checked_record.record.title}</a>'
            ])
        return JsonResponse({'data': data})


class Dashboard(View):
    name = 'records/dashboard.html'

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_roles(roles=['ktto', 'rdco', 'itso', 'tbi']))
    def get(self, request):
        return render(request, self.name)

    def post(self, request):
        if is_ajax(request=request):
            checked_records = CheckedRecord.objects.filter(status='approved', checked_by__in=Subquery(
                User.objects.filter(role=5).values('pk')))
            records = Record.objects.filter(pk__in=Subquery(checked_records.values('record_id')))
            # graphs
            if request.POST.get('graphs'):
                basic_count = records.filter(classification=1).count()
                applied_count = records.filter(classification=2).count()
                psced_count = []
                records_per_year_count = []
                psced_per_year_count = []
                adviser_pending_count = []
                psced_classifications = PSCEDClassification.objects.all()
                records_per_year = records.values('year_accomplished').annotate(
                    year_count=Count('year_accomplished')).order_by('year_accomplished')[:10]
                for psced in psced_classifications:
                    psced_count.append({'name': psced.name, 'count': records.filter(
                        psced_classification=PSCEDClassification.objects.get(pk=psced.id)).count()})
                for record_per_year in records_per_year:
                    records_per_year_count.append(
                        {'year': record_per_year['year_accomplished'], 'count': record_per_year['year_count']})

                # NUMBER OF CLASSIFICATIONS PER YEAR
                # year_count = records.values('year_accomplished').distinct().annotate(psced_count=Count('psced_classification', distinct=True)).order_by('year_accomplished')
                # Convert all year_count values into a list of tuples
                # tuples_year_count = [tuple(d.values()) for d in year_count]
                # for row in tuples_year_count:
                #    psced_per_year_count.append({'year': row[0], 'psced_count': row[1]})
                with connection.cursor() as cursor:
                    cursor.execute(
                        "SELECT year_accomplished, COUNT(year_accomplished) AS year_count FROM (SELECT DISTINCT year_accomplished, psced_classification_id FROM (select year_accomplished, psced_classification_id from records_record inner join records_checkedrecord on records_record.id=record_id inner join accounts_user on records_checkedrecord.checked_by_id=accounts_user.id where accounts_user.role_id=5 and records_checkedrecord.status='approved') as recordtbl) as tbl GROUP BY year_accomplished")
                    rows = cursor.fetchall()
                    print(rows)
                    for row in rows:
                        psced_per_year_count.append({'year': row[0], 'psced_count': row[1]})

                # Pending Adviser
                # adviser_exclude = CheckedRecord.objects.select_related('record').all()
                # adviser_pending = Record.objects.exclude(pk__in=Subquery(adviser_exclude.values('record').distinct())).values('pk', 'title')
                # adviser_pending_count = adviser_pending.count()
                with connection.cursor() as cursor:
                    cursor.execute(
                        f"select records_record.id, records_record.title, records_checkedrecord.checked_by_id from records_record left join records_checkedrecord on records_record.id = records_checkedrecord.record_id where checked_by_id is null")
                    rows = cursor.fetchall()
                    print(rows)
                    adviser_pending_count = len(rows)

                # Pending KTTO
                # ktto_exclude = CheckedRecord.objects.select_related('record').filter(Q(checked_by__in=Subquery(User.objects.filter(role=4).values('pk'))) | Q(checked_by__in=Subquery(User.objects.filter(role=7).values('pk'))))
                # ktto_include = CheckedRecord.objects.select_related('record').filter(status='approved', checked_by__in=Subquery(User.objects.filter(role=3).values('pk')))
                # ktto_pending = Record.objects.filter(pk__in=Subquery(ktto_include.values('record'))).exclude(pk__in=Subquery(ktto_exclude.values('record'))).values('pk', 'title')
                # ktto_pending_count = ktto_pending.count()
                with connection.cursor() as cursor:
                    cursor.execute(
                        "SELECT records_record.id, records_record.title FROM records_record INNER JOIN records_checkedrecord ON records_record.id = records_checkedrecord.record_id INNER JOIN accounts_user ON records_checkedrecord.checked_by_id = accounts_user.id WHERE accounts_user.role_id = 3 AND records_checkedrecord.status = 'approved' AND records_record.id NOT IN (SELECT records_checkedrecord.record_id FROM records_checkedrecord INNER JOIN accounts_user ON records_checkedrecord.checked_by_id = accounts_user.id WHERE accounts_user.role_id = 4 or accounts_user.role_id = 7)")
                    rows = cursor.fetchall()
                    print(rows)
                    ktto_pending_count = len(rows)

                # Pending RDCO
                # rdco_exclude = CheckedRecord.objects.select_related('record').filter(checked_by=Subquery(User.objects.filter(role=5).values('pk')))
                # rdco_include = CheckedRecord.objects.select_related('record').filter(checked_by__in=Subquery(User.objects.filter(role=4).values('pk')), status='approved')
                # rdco_pending = Record.objects.filter(pk__in=Subquery(rdco_include.values('record'))).exclude(pk__in=Subquery(rdco_exclude.values('record'))).values('pk', 'title')
                # print(rdco_pending)
                # rdco_pending_count = rdco_pending.count()
                with connection.cursor() as cursor:
                    cursor.execute(
                        "SELECT records_record.id, records_record.title FROM records_record INNER JOIN records_checkedrecord ON records_record.id = records_checkedrecord.record_id INNER JOIN accounts_user ON records_checkedrecord.checked_by_id = accounts_user.id WHERE accounts_user.role_id = 4 AND records_checkedrecord.status = 'approved' AND records_record.id NOT IN (SELECT records_checkedrecord.record_id FROM records_checkedrecord INNER JOIN accounts_user ON records_checkedrecord.checked_by_id = accounts_user.id WHERE accounts_user.role_id = 5)")
                    rows = cursor.fetchall()
                    rdco_pending_count = len(rows)
                    print(rows)
                record_uploads = RecordUpload.objects.all()

                def get_doc_counts(docs):
                    patent_count = 0
                    utility_model_count = 0
                    industrial_design_count = 0
                    trademark_count = 0
                    copyright_count = 0
                    for data in docs:
                        if data.upload.name == 'Patent':
                            patent_count += 1
                        if data.upload.name == 'Utility Model':
                            utility_model_count += 1
                        if data.upload.name == 'Industrial Design':
                            industrial_design_count += 1
                        if data.upload.name == 'Trademark':
                            trademark_count += 1
                        if data.upload.name == 'Copyright':
                            copyright_count += 1
                    return [patent_count, utility_model_count, industrial_design_count, trademark_count,
                            copyright_count]

                for_application = get_doc_counts(
                    record_uploads.filter(record_upload_status=RecordUploadStatus.objects.get(pk=1)))
                reviewed = get_doc_counts(
                    record_uploads.filter(record_upload_status=RecordUploadStatus.objects.get(pk=2)))
                filed = get_doc_counts(record_uploads.filter(record_upload_status=RecordUploadStatus.objects.get(pk=3)))
                approved = get_doc_counts(
                    record_uploads.filter(record_upload_status=RecordUploadStatus.objects.get(pk=4)))
                disapproved = get_doc_counts(
                    record_uploads.filter(record_upload_status=RecordUploadStatus.objects.get(pk=5)))
                return JsonResponse({'success': True, 'basic': basic_count, 'applied': applied_count,
                                     'psced_count': psced_count, 'records_per_year_count': records_per_year_count,
                                     'psced_per_year_count': psced_per_year_count,
                                     'adviser_pending_count': adviser_pending_count,
                                     'ktto_pending_count': ktto_pending_count, 'rdco_pending_count': rdco_pending_count,
                                     'for_application': for_application, 'reviewed': reviewed, 'filed': filed,
                                     'approved': approved, 'disapproved': disapproved})


class LogsView(View):
    name = 'records/dashboard/logs_view.html'

    @method_decorator(authorized_roles(roles=['ktto', 'rdco', 'tbi', 'itso']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        return render(request, self.name)

    def post(self, request):
        if is_ajax(request=request):
            data = []
            logs = Log.objects.all()
            for log in logs:
                data.append([log.pk, log.action, log.user.username, log.date_created.strftime("%Y-%m-%d %H:%M:%S")])
            return JsonResponse({'data': data})


class DashboardLogsRecordView(View):
    name = 'records/dashboard/logs_view_record.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
        'is_owner': True,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_record_user())
    def get(self, request, record_id):
        owners = UserRecord.objects.filter(record=Record.objects.get(pk=record_id))
        self.context['owners'] = owners
        checked_records = CheckedRecord.objects.filter(record=Record.objects.get(pk=record_id))
        adviser_checked = {'status': 'pending'}
        ktto_checked = {'status': 'pending'}
        rdco_checked = {'status': 'pending'}
        role_checked = False
        record = Record.objects.get(pk=record_id)
        research_record = ResearchRecord.objects.filter(Q(proposal=record) | Q(research=record)).first()
        for checked_record in checked_records:
            if checked_record.checked_by.role.id == 3:
                adviser_checked = {'status': checked_record.status, 'content': checked_record}
            if checked_record.checked_by.role.id == 4 or checked_record.checked_by.role.id == 7:
                ktto_checked = {'status': checked_record.status, 'content': checked_record}
            if checked_record.checked_by.role.id == 5:
                rdco_checked = {'status': checked_record.status, 'content': checked_record}
            if checked_record.checked_by.role.id == request.user.role.pk:
                role_checked = True
        self.context['adviser_checked'] = adviser_checked
        self.context['ktto_checked'] = ktto_checked
        self.context['rdco_checked'] = rdco_checked
        self.context['role_checked'] = role_checked
        self.context['record'] = record
        self.context['is_removable'] = True
        self.context['research_record'] = research_record
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # removing record
            if request.POST.get('remove', 'false') == 'true':
                del_record = Record.objects.get(pk=record_id)
                del_record.abstract_file.delete()
                del_record_uploads = RecordUpload.objects.filter(record=del_record)
                for del_record_upload in del_record_uploads:
                    del_record_upload.file.delete()
                del_record.delete()
                return JsonResponse({'success': True})
            # updating record tags
            elif request.POST.get('tags_update', 'false') == 'true':
                return JsonResponse(update_record_tags(request, record_id))
            # get uploaded document data
            elif request.POST.get('get_document', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                checked_uploads = CheckedUpload.objects.filter(record_upload=record_upload).order_by('-date_checked')
                comments = []
                checked_bys = []
                checked_dates = []
                for checked_upload in checked_uploads:
                    comments.append(checked_upload.comment)
                    checked_bys.append(checked_upload.checked_by.username)
                    checked_dates.append(checked_upload.date_checked)
                if record_upload is None:
                    return JsonResponse({'success': False, 'doc-title': upload.name})
                else:
                    return JsonResponse({'success': True,
                                         'doc-title': record_upload.upload.name,
                                         'doc-status': record_upload.record_upload_status.name,
                                         'is-ip': record_upload.is_ip,
                                         'for-commercialization': record_upload.for_commercialization,
                                         'comments': comments,
                                         'checked_bys': checked_bys,
                                         'checked_dates': checked_dates,
                                         'record-upload-id': record_upload.pk})
            # POSTING COMMENTS
            elif request.POST.get('post_comment', 'false') == 'true':
                upload = Upload.objects.get(pk=request.POST.get('upload_id', 0))
                record = Record.objects.get(pk=request.POST.get('record_id', 0))
                comment = request.POST.get('comment', '')
                record_upload = RecordUpload.objects.filter(upload=upload, record=record).first()
                CheckedUpload(comment=comment, checked_by=request.user, record_upload=record_upload).save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False})



# dashboard manage records table
class ViewManageRecords(View):
    name = 'records/dashboard/manage_records.html'

    @method_decorator(authorized_roles(roles=['adviser', 'ktto', 'rdco', 'tbi', 'itso']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        return render(request, self.name)

    def post(self, request):
        if is_ajax(request=request):
            records = Record.objects.all()
            data = []

            # filtering records
            if request.POST.get('is-filter', '0') == '1':
                is_ip = request.POST.get('is-ip', 0)
                commercialization = request.POST.get('for-commercialization', 0)
                community_ext = request.POST.get('community-ext', 0)
                no_tags = request.POST.get('no-tags', 0)
                record_publish_status = request.POST.get('record-publish-status', 0)
                if record_publish_status == '1':
                    checked_records = CheckedRecord.objects.filter(status='approved', checked_by__in=Subquery(
                        User.objects.filter(role=5).values('pk')))
                    records = records.filter(pk__in=Subquery(checked_records.values('record_id')))
                elif record_publish_status == '2':
                    checked_records = CheckedRecord.objects.filter(status='approved', checked_by__in=Subquery(
                        User.objects.filter(role=5).values('pk')))
                    records = records.exclude(pk__in=Subquery(checked_records.values('record_id')))
                if is_ip == '1':
                    records = records.filter(is_ip=True)
                if commercialization == '1':
                    records = records.filter(for_commercialization=True)
                if community_ext == '1':
                    records = records.filter(community_extension=True)
                if no_tags == '1':
                    records = records.filter(community_extension=False, is_ip=False, for_commercialization=False)

            # removing records
            elif request.POST.get('remove'):
                titles = request.POST.getlist('titles[]')
                for title_id in titles:
                    del_record = Record.objects.get(pk=int(title_id))
                    del_record.abstract_file.delete()
                    del_record.delete()
                return JsonResponse({'success': True})

            # Edit record code
            elif request.POST.get('edit-record-code', 'false') == 'true':
                record = Record.objects.get(pk=request.POST.get('record-id', None))
                record_code = request.POST.get('record-code', None)
                record.code = record_code
                record.save()

            for record in records:
                tags = ''
                if record.is_ip:
                    tags = f'<div class="badge badge-primary">IP</div>&nbsp;'
                if record.for_commercialization:
                    tags += f'<div class="badge badge-success">For commercialization</div>&nbsp;'
                if record.community_extension:
                    tags += f'<div class="badge badge-secondary">Community extension</div>&nbsp;'
                data.append([
                    '',
                    record.pk,
                    f'{record.code} <a href="#" onclick="editCode({record.pk},\'{record.code}\')"><i class="fa fa-pen fa-md"></i></a>',
                    f'<a href="/dashboard/manage/records/{record.pk}">{record.title}</a>',
                    record.record_type.name,
                    record.date_created.strftime("%Y-%m-%d %H:%M:%S"),
                    tags,
                ])

            return JsonResponse({'data': data})


# Dashboard manage record template
class DashboardManageRecord(View):
    name = 'records/dashboard/manage_view_record.html'
    author_roles = AuthorRole.objects.all()
    classifications = Classification.objects.all()
    psced_classifications = PSCEDClassification.objects.all().order_by('name')
    conference_levels = ConferenceLevel.objects.all()
    budget_types = BudgetType.objects.all()
    collaboration_types = CollaborationType.objects.all()
    publication_levels = PublicationLevel.objects.all()
    uploads = Upload.objects.all()
    checked_record_form = CheckedRecordCommentForm()
    context = {
        'author_roles': author_roles,
        'classifications': classifications,
        'psced_classifications': psced_classifications,
        'conference_levels': conference_levels,
        'budget_types': budget_types,
        'collaboration_types': collaboration_types,
        'publication_levels': publication_levels,
        'uploads': uploads,
        'checked_record_form': checked_record_form,
        'is_owner': True,
    }

    @method_decorator(login_required(login_url='/'))
    @method_decorator(authorized_record_user())
    def get(self, request, record_id):
        view_record_controller = ViewRecordController()
        data = view_record_controller.getRecordInfo(request, record_id)
        self.context.update(data)
        return render(request, self.name, self.context)

    def post(self, request, record_id):
        if is_ajax(request=request):
            # removing record
            if request.POST.get('remove', 'false') == 'true':
                del_record = Record.objects.get(pk=record_id)
                del_record.abstract_file.delete()
                del_record_uploads = RecordUpload.objects.filter(record=del_record)
                for del_record_upload in del_record_uploads:
                    del_record_upload.file.delete()
                del_record.delete()
                return JsonResponse({'success': True})
        else:
            # approving or declining record
            evaluation_controller = EvaluationController()
            evaluation_controller.evaluateRecord(record_id)
            
            


class DashboardManageAccounts(View):
    name = 'records/dashboard/accounts_view.html'

    @method_decorator(authorized_roles(roles=['ktto', 'rdco', 'tbi', 'itso']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        user_roles = UserRole.objects.all()
        context = {
            'user_roles': user_roles,
        }
        return render(request, self.name, context)

    def post(self, request):
        if is_ajax(request=request):
            # removing accounts
            if request.POST.get('remove-accounts'):
                accounts = request.POST.getlist('accounts[]')
                success = False
                for account_id in accounts:
                    del_account = User.objects.get(pk=int(account_id))
                    if not del_account.is_superuser:
                        del_account.delete()
                        success = True
                return JsonResponse({'success': success})
            # accounts role change
            elif request.POST.get('role-change') == 'true':
                accounts = request.POST.getlist('accounts[]')
                accounts_str = ''

                role_id = int(request.POST.get('role-radio'))
                for account_id in accounts:
                    user = User.objects.get(pk=int(account_id))
                    user.role = UserRole.objects.get(pk=role_id)
                    user.save()
                    if account_id == accounts[0]:
                        accounts_str += user.username
                    else:
                        accounts_str += f', {user.username}'
                    RoleRequest.objects.filter(user=user).delete()
                Log(user=request.user,
                    action=f'accounts: {accounts_str} account_role changed to \"{user.role}\" by: {request.user.username}').save()
                # roleRequestNotify(request.user.id, user.id)
                roleRequestApproved(request, request.user.id, user.id)


class LockoutPage(View):
    name = "records/lockout_page.html"

    def get(self, request):
        return render(request, self.name)


class LockedAccountsView(View):
    name = "records/dashboard/reset_locked_accounts.html"

    @method_decorator(authorized_roles(roles=['ktto', 'rdco', 'tbi', 'itso']))
    @method_decorator(login_required(login_url='/'))
    def get(self, request):
        attempts = AccessAttempt.objects.all()
        context = {
            'attempts': attempts,
        }
        return render(request, self.name, context)

    def post(self, request):
        if request.method == 'POST':
            if request.POST.get('resetAccount'):
                accounts = request.POST.getlist('listOfAccounts[]')
                print(accounts)
                for account in accounts:
                    reset(username=account)
                return JsonResponse({'success': True})


def get_all_download_requests(request):
    if request.method == 'POST':
        data = []
        download_requests = RecordDownloadRequest.objects.filter(is_marked=False)
        for record in download_requests:
            data.append([
                record.pk,
                '',
                record.record.title,
                record.sent_by.username,
                f'<div><a href="#" data-toggle="modal" data-target="#approvedDownloadModal" class="{record.sent_by.username}-{record.pk}" id="approve-btn-{record.record.pk}" onclick="return getModalId(this.id, this.className);"> Approve </a></div>'
            ])

        return JsonResponse({'data': data})


def approved_download_requests(request):
    if request.method == 'POST':
        if 'approvedDownloadBtn' in request.POST:
            record_id = request.POST.get('record_number')
            username = request.POST.get('username')
            request_id = request.POST.get('request_id')
            user = User.objects.get(username=username)
            download_request = RecordDownloadRequest.objects.get(pk=request_id)
            download_request.is_marked = True
            download_request.save()

            approvedDownloadRequest(request, request.user.id, record_id, user.id)
            messages.success(request, 'Download request approved')
            return redirect("records-requests")
        else:
            messages.error(request, 'Error encountered while approving request')
            return redirect("records-requests")


@require_POST
def approve_attachments(request, checked_upload_id):
    try:
        checked_upload = CheckedUpload.objects.get(id=checked_upload_id)

        # Extract comments from POST data
        comment = request.POST.get('comment', '')

        # Update the comment in CheckedUpload
        checked_upload.comment = comment
        checked_upload.save()

        # Your existing logic for approving the record can go here
        if checked_upload:  # Replace with your actual condition
            checked_upload.is_approved = True
            checked_upload.save()

            # Update the is_approved field in RecordUpload to 1
            record_upload = checked_upload.record_upload
            record_upload.is_approved = True
            record_upload.save()

            # Additional logic if needed when it is approved

            return JsonResponse({'success': True, 'message': 'Record approved successfully'})
        else:
            return JsonResponse({'success': True, 'message': 'Record checked but not approved'})

    except CheckedUpload.DoesNotExist as e:
        return JsonResponse({'success': False, 'message': 'CheckedUpload does not exist'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})


@require_POST
def decline_attachments(request, checked_upload_id):
    try:
        checked_upload = CheckedUpload.objects.get(id=checked_upload_id)

        # Extract comments from POST data
        comment = request.POST.get('comment', '')

        # Update the comment in CheckedUpload
        checked_upload.comment = comment
        checked_upload.save()

        # Your existing logic for approving the record can go here
        if checked_upload:  # Replace with your actual condition
            checked_upload.is_declined = True
            checked_upload.save()

            # Update the is_approved field in RecordUpload to 1
            record_upload = checked_upload.record_upload
            record_upload.is_declined = True
            record_upload.save()

            # Additional logic if needed when it is approved

            return JsonResponse({'success': True, 'message': 'Record approved successfully'})
        else:
            return JsonResponse({'success': True, 'message': 'Record checked but not approved'})

    except CheckedUpload.DoesNotExist as e:
        return JsonResponse({'success': False, 'message': 'CheckedUpload does not exist'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

    @require_POST
    def approve_attachments(request, record_upload_id):
        record_upload = get_object_or_404(RecordUpload, id=record_upload_id)
        comment = request.POST.get('comment', '')
        # Perform approval logic here (e.g., update the 'is_approved' field)
        record_upload.is_approved = True
        record_upload.save()
        return JsonResponse({'status': 'success'})

    @require_POST
    def decline_attachments(request, record_upload_id):
        record_upload = get_object_or_404(RecordUpload, id=record_upload_id)
        comment = request.POST.get('comment', '')
        # Perform decline logic here (e.g., update the 'is_approved' field)
        record_upload.is_approved = False
        record_upload.save()
        return JsonResponse({'status': 'success'})


def calculate_total_docx_pages(docx_path):
    try:
        doc = Document(docx_path)
        total_pages = len(doc.element.xpath('//w:sectPr'))

        return total_pages
    except Exception as e:
        # Handle the exception or log it as needed
        print(f"Error calculating total DOCX pages: {e}")
        return 0
def read_docx_content(file_path):
    """
    Function to read text content from a DOCX file.
    """
    doc = Document(file_path)
    text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text_content

def extract_images_from_docx(file_path):
    """
    Function to extract images from a DOCX file.
    """
    doc = Document(file_path)
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = base64.b64encode(rel.target_part.blob).decode('utf-8')
            images.append({'data': image_data})
    return images

def view_file_content(request, record_upload_id):
    try:
        # Retrieve the RecordUpload object or return a 404 response if not found
        record_upload = get_object_or_404(RecordUpload, pk=record_upload_id)

        # Initialize variables
        content = None
        text_content = None
        images = None
        total_pages = ''

        # Open the file and convert its content to base64
        with open(record_upload.file.path, 'rb') as file:
            if record_upload.file.name.endswith('.pdf'):
                # For PDF files
                content = base64.b64encode(file.read()).decode('utf-8')
            elif record_upload.file.name.endswith('.docx'):
                # For DOCX files
                text_content = read_docx_content(record_upload.file.path)
                images = extract_images_from_docx(record_upload.file.path)
                total_pages = calculate_total_docx_pages(record_upload.file.path)
            else:
                raise ValueError(f'Unsupported file type: {record_upload.file.name}')

        # Return the content as JSON
        return JsonResponse({
            'content': content,
            'text_content': text_content,
            'images': images,
            'total_pages': total_pages,  # Include total pages for DOCX files
            'fileType': get_file_type(record_upload.file.name)
        })
    except Exception as e:
        # Log the exception
        traceback.print_exc()

        # Return an error response
        return JsonResponse({'error': f'Error fetching file content: {str(e)}'}, status=500)
def read_docx_content(file):
    document = Document(file)
    content = [paragraph.text for paragraph in document.paragraphs]
    return '\n'.join(content)

def get_file_type(file_name):
    """
    Function to determine the file type based on its extension.
    """
    if file_name.lower().endswith('.pdf'):
        return 'pdf'
    elif file_name.lower().endswith('.docx'):
        return 'docx'
    else:
        return 'unknown'


def download_docx_file(request, record_upload_id):
    record_upload = get_object_or_404(RecordUpload, pk=record_upload_id)

    # Assuming you have a field named 'file' in your model
    docx_file = record_upload.file

    # Check if the file is a DOCX file before proceeding
    if not docx_file.name.endswith('.docx'):
        # Return a response indicating that only DOCX files are allowed
        return HttpResponse()

    # Open and read the content of the docx file
    with docx_file.open(mode='rb') as f:
        docx_content = f.read()

    # Set the response content type
    content_type, encoding = mimetypes.guess_type(docx_file.name)
    response = HttpResponse(content_type=content_type)

    # Set the content-disposition for the response
    response['Content-Disposition'] = f'attachment; filename={docx_file.name}'

    # Write the docx content to the response
    response.write(docx_content)

    return response



# views.py

from django.shortcuts import redirect
import requests
import base64
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime, timedelta
from .models import Subscription
import logging

logger = logging.getLogger(__name__)

paymongo_api_key = 'sk_test_PUL9xuAM8Sm9GLh3FGura1vr'  # Replace this with your actual Paymongo API key

def create_payment_link_view(request):
    try:
        tier = request.GET.get('tier')  # Extract the 'tier' parameter from the URL query parameters

        if tier == 'standard':
            amount = 10000  # 100.00 pesos in cents
        elif tier == 'premium':
            amount = 14900  # 149.00 pesos in cents
        else:
            # Handle invalid tier
            return redirect('/')  # Redirect to homepage or error page

        url = 'https://api.paymongo.com/v1/links'
        payload = {
            'data': {
                'attributes': {
                    'amount': amount,
                    'description': f'IPAMS {tier.capitalize()} Subscription Payment',
                    'remarks': 'pay',
                    'status': 'unpaid'
                }
            }
        }
        encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
        headers = {
            'Authorization': f'Basic {encoded_api_key}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()

        data = response.json()
        checkout_url = data['data']['attributes']['checkout_url']
        stored_link_id = data['data']['id']
        
        # Store the link_id in the session to verify payment later
        request.session['stored_link_id'] = stored_link_id

        return redirect(checkout_url)

    except Exception as error:
        logger.error('Error creating payment link: %s', error)
        return redirect('/')  # Redirect to homepage or error page

    
logger = logging.getLogger(__name__)

def create_payment_link_view(request):
    try:
        tier = request.GET.get('tier')  # Extract the 'tier' parameter from the URL query parameters

        # Fetch the SubscriptionPlan based on the tier
        plan = get_object_or_404(SubscriptionPlan, plan_name__iexact=tier)
        amount = int(plan.price * 100)  # Convert to cents

        url = 'https://api.paymongo.com/v1/links'
        payload = {
            'data': {
                'attributes': {
                    'amount': amount,
                    'description': f'IPAMS {tier.capitalize()} Subscription Payment',
                    'remarks': 'pay',
                    'status': 'unpaid'
                }
            }
        }
        encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
        headers = {
            'Authorization': f'Basic {encoded_api_key}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, json=payload, headers=headers)
        response.raise_for_status()

        data = response.json()
        checkout_url = data['data']['attributes']['checkout_url']
        stored_link_id = data['data']['id']
        
        # Store the link_id in the session to verify payment later
        request.session['stored_link_id'] = stored_link_id

        return redirect(checkout_url)

    except Exception as error:
        logger.error('Error creating payment link: %s', error)
        return redirect('/')  # Redirect to homepage or error page

import requests
from django.http import JsonResponse
import base64

@login_required
def renew_subscription(request):
    if request.method == 'POST':
        try:
            # Get the user's current subscription
            subscription = Subscription.objects.get(user_id=request.user, status='active')

            # Check the plan_id from the subscription and determine the tier and amount
            plan_id = subscription.plan_id_id
            if plan_id == 1:
                tier = 'free trial'
                amount = 0  # Free trial, no charge
            elif plan_id == 2:
                tier = 'standard'
                amount = 15000  # 150 PHP in cents
            elif plan_id == 3:
                tier = 'premium'
                amount = 20000  # 200 PHP in cents
            else:
                return JsonResponse({'success': False, 'message': 'Invalid subscription plan.'})

            # Create a PayMongo payment link from the backend (server-to-server)
            url = 'https://api.paymongo.com/v1/links'
            payload = {
                'data': {
                    'attributes': {
                        'amount': amount,
                        'description': f'{tier.capitalize()} Subscription Payment',
                        'remarks': 'Subscription Payment',
                        'status': 'unpaid'
                    }
                }
            }

            encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
            headers = {
                'Authorization': f'Basic {encoded_api_key}',
                'Content-Type': 'application/json'
            }

            # Make the API request to PayMongo from the server
            response = requests.post(url, json=payload, headers=headers)
            response.raise_for_status()

            # Extract the payment URL and return it as JSON response
            data = response.json()
            checkout_url = data['data']['attributes']['checkout_url']

            return JsonResponse({'success': True, 'checkout_url': checkout_url})

        except Subscription.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Active subscription not found.'})
        except requests.exceptions.HTTPError as http_err:
            return JsonResponse({'success': False, 'message': f'Payment processing error: {http_err}'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': 'An unexpected error occurred.'})



def get_payment_link_and_check_status(link_id):
    try:
        url = f"https://api.paymongo.com/v1/links/{link_id}"
        encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
        headers = {
            "accept": "application/json",
            "authorization": f"Basic {encoded_api_key}"
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # Raise an error for non-2xx status codes

        data = response.json()
        if data['data']:
            status = data['data']['attributes']['status']
            return status
        else:
            return None  # If no data is returned

    except Exception as error:
        logger.error('Error getting payment link or checking status: %s', error)
        return None  # Return None if an error occurs


from django.db import connections
from django.utils.timezone import now

def verify_subscription(request):
    if request.method == 'POST':
        reference_number = request.POST.get('reference_number')

        if not reference_number:
            logger.error('Reference number is required.')
            return JsonResponse({'success': False, 'message': 'Reference number is required.'})

        url = f"https://api.paymongo.com/v1/links/{reference_number}"
        encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
        headers = {
            "accept": "application/json",
            "authorization": f"Basic {encoded_api_key}"
        }

        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()

            if data and 'data' in data:
                payment_data = data['data']
                attributes = payment_data.get('attributes', {})
                logger.info(f'Payment data attributes: {attributes}')

                if attributes.get('status') == 'paid':
                    user = request.user

                    # Delete any existing subscriptions for the user
                    Subscription.objects.filter(user_id=user).delete()

                    # Determine the subscription plan based on the amount paid
                    amount_paid = attributes.get('amount') / 100  # Convert cents to pesos

                    try:
                        plan = SubscriptionPlan.objects.get(price=amount_paid)
                    except SubscriptionPlan.DoesNotExist:
                        logger.error('No subscription plan matches the amount paid.')
                        return JsonResponse({'success': False, 'message': 'No subscription plan matches the amount paid.'})

                    # Create a new subscription
                    subscription = Subscription.objects.create(
                        start_date=datetime.now().date(),
                        end_date=(datetime.now() + timedelta(days=180)).date(),
                        user_id=user,
                        plan_id=plan,
                        status='active'
                    )

                    # Update user's subscription status in the main database
                    user.is_subscribed = True
                    user.subscription_status = 'paid'
                    user.sub_id = subscription.sub_id
                    user.save()

                    # Determine the subscription type for nalc
                    if plan.price == 0:
                        nalc_subscription_type = 'FREE TRIAL'
                    elif plan.price == 150:
                        nalc_subscription_type = 'STANDARD'
                    elif plan.price == 200:
                        nalc_subscription_type = 'PREMIUM'
                    else:
                        logger.error('Unknown subscription plan price for nalc integration.')
                        return JsonResponse({'success': False, 'message': 'Unknown subscription plan price for nalc integration.'})

                    # Insert or update subscription field in nalc database
                    try:
                        logger.info(f"Updating NALC subscription for user {user.email} with plan {nalc_subscription_type}")
                        
                        # Fetch the user_id from the nalc database using the email
                        with connections['nalc'].cursor() as cursor:
                            cursor.execute("SELECT id FROM backend_user WHERE email = %s", [user.email])
                            nalc_user_id_result = cursor.fetchone()
                            
                        if nalc_user_id_result:
                            nalc_user_id = nalc_user_id_result[0]
                            logger.info(f"Found NALC user with id {nalc_user_id} for email {user.email}")
                            
                            # Now update the subscription field using the nalc_user_id
                            with connections['nalc'].cursor() as cursor:
                                cursor.execute("""
                                    UPDATE backend_user
                                    SET subscription = %s
                                    WHERE id = %s;
                                """, [nalc_subscription_type, nalc_user_id])

                            logger.info(f'Successfully updated NALC user subscription to {nalc_subscription_type}')
                        else:
                            logger.error(f"User with email {user.email} not found in NALC database. No record updated.")
                            return JsonResponse({'success': False, 'message': 'User not found in NALC database by email.'})

                    except Exception as e:
                        logger.error(f'Error updating NALC database: {e}')
                        return JsonResponse({'success': False, 'message': 'Failed to update NALC subscription.'})




                    return JsonResponse({'success': True, 'message': 'Subscription verified and updated successfully.'})
                else:
                    logger.error('Payment for the reference number is not yet completed.')
                    return JsonResponse({'success': False, 'message': 'Payment for the reference number is not yet completed.'})
            else:
                logger.error('Invalid reference number.')
                return JsonResponse({'success': False, 'message': 'Invalid reference number.'})
        except requests.HTTPError as e:
            if response.status_code == 404:
                logger.error('Payment link not found. Please check the reference number.')
                return JsonResponse({'success': False, 'message': 'Payment link not found. Please check the reference number.'})
            else:
                logger.error(f'Error in Paymongo API request: {e}')
                return JsonResponse({'success': False, 'message': f'Error in Paymongo API request: {e}'})
        except requests.RequestException as e:
            logger.error(f'Error in Paymongo API request: {e}')
            return JsonResponse({'success': False, 'message': f'Error in Paymongo API request: {e}'})
        except Exception as e:
            logger.error(f'Unexpected error: {e}')
            return JsonResponse({'success': False, 'message': f'Unexpected error: {e}'})
    else:
        logger.error('Invalid request method.')
        return JsonResponse({'success': False, 'message': 'Invalid request method.'})

def verify_renewal_subscription(request):
    if request.method == 'POST':
        reference_number = request.POST.get('reference_number')

        if not reference_number:
            logger.error('Reference number is required.')
            return JsonResponse({'success': False, 'message': 'Reference number is required.'})

        url = f"https://api.paymongo.com/v1/links/{reference_number}"
        encoded_api_key = base64.b64encode(paymongo_api_key.encode()).decode()
        headers = {
            "accept": "application/json",
            "authorization": f"Basic {encoded_api_key}"
        }

        try:
            response = requests.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()

            if data and 'data' in data:
                payment_data = data['data']
                attributes = payment_data.get('attributes', {})
                logger.info(f'Payment data attributes: {attributes}')

                if attributes.get('status') == 'paid':
                    user = request.user

                    # Access the user's current active subscription
                    try:
                        subscription = Subscription.objects.get(user_id=user, status='active')
                    except Subscription.DoesNotExist:
                        logger.error('Active subscription not found for user.')
                        return JsonResponse({'success': False, 'message': 'No active subscription found to renew.'})

                    # Extend the end date by 6 months
                    new_end_date = subscription.end_date + timedelta(days=180)
                    subscription.end_date = new_end_date
                    subscription.save()

                    user.subscription_status = 'renewed'
                    user.save()

                    logger.info(f'Subscription for user {user} extended. New end date: {new_end_date}')
                    return JsonResponse({'success': True, 'message': 'Subscription renewed successfully.'})
                else:
                    logger.error('Payment for the reference number is not yet completed.')
                    return JsonResponse({'success': False, 'message': 'Payment for the reference number is not yet completed.'})
            else:
                logger.error('Invalid reference number.')
                return JsonResponse({'success': False, 'message': 'Invalid reference number.'})
        except requests.HTTPError as e:
            if response.status_code == 404:
                logger.error('Payment link not found. Please check the reference number.')
                return JsonResponse({'success': False, 'message': 'Payment link not found. Please check the reference number.'})
            else:
                logger.error(f'Error in Paymongo API request: {e}')
                return JsonResponse({'success': False, 'message': f'Error in Paymongo API request: {e}'})
        except requests.RequestException as e:
            logger.error(f'Error in Paymongo API request: {e}')
            return JsonResponse({'success': False, 'message': f'Error in Paymongo API request: {e}'})
        except Exception as e:
            logger.error(f'Unexpected error: {e}')
            return JsonResponse({'success': False, 'message': f'Unexpected error: {e}'})
    else:
        logger.error('Invalid request method.')
        return JsonResponse({'success': False, 'message': 'Invalid request method.'})
    
def home_view(request):
    user = request.user
    show_modal = False

    if request.user.is_authenticated:
        try:
            subscription = Subscription.objects.get(user_id=user, status='active')
            days_left = (subscription.end_date - datetime.now().date()).days
            if days_left <= 7:
                show_modal = True
        except Subscription.DoesNotExist:
            pass

    return render(request, 'base.html', {'show_modal': show_modal})



@login_required
def cancel_subscription(request):
    try:
        subscription = Subscription.objects.get(user_id=request.user, status='active')
        subscription.status = 'inactive'
        subscription.save()
        request.user.is_subscribed = False
        request.user.subscription_status = 'unpaid'
        request.user.save()
        return JsonResponse({'success': True, 'message': 'Subscription canceled successfully.'})
    except Subscription.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Active subscription not found.'})
    

from django.http import JsonResponse

from django.contrib.auth.decorators import login_required

@login_required
def check_verification_status(request):
    user = request.user
    is_subscribed = user.is_subscribed
    is_verified = (is_subscribed == 1)  
    return JsonResponse({'is_verified': is_verified})

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseForbidden
from .models import SubscriptionPlan

from django.http import JsonResponse

@login_required
def update_price(request):
    if request.user.role_id != 7:
        return HttpResponseForbidden("You do not have permission to edit this.")

    if request.method == 'POST':
        tier = request.POST.get('tier')
        price = request.POST.get('price')

        print(f"Updating {tier} plan with new price: {price}")
        
        plan = get_object_or_404(SubscriptionPlan, plan_name=tier)
        plan.price = price
        plan.save()

       
        return JsonResponse({"message": "updated"})
    
    return HttpResponseForbidden("Invalid request method.")


def subscription_plans_view(request):
    try:
        plans = SubscriptionPlan.objects.all()
        print("Plans retrieved: ", plans)  # Debugging statement
    except Exception as e:
        print("An error occurred: ", e)
    return render(request, 'ipams/subscribe.html', {'plans': plans})


from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.db import connections
from django.utils.timezone import now
import json
import traceback

@csrf_exempt
def update_api_key(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            api_key = data.get('api_key')
            if not api_key:
                return JsonResponse({'success': False, 'error': 'API Key is required'}, status=400)

            with connections['nalc'].cursor() as cursor:
                cursor.execute("""
                    INSERT INTO backend_openai_api (api_key, created_at)
                    VALUES (%s, %s)
                    ON DUPLICATE KEY UPDATE
                        created_at = VALUES(created_at);
                """, [api_key, now()])

            return JsonResponse({'success': True})
        except Exception as e:
            print(f"Error updating API key: {e}")
            print(traceback.format_exc())
            return JsonResponse({'success': False, 'error': 'An error occurred while updating the API key.'}, status=500)
    return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=405)


# records/views.py
from django.shortcuts import render
from .models import SubscriptionPlan

def manage_subscriptions(request):
    subscriptions = SubscriptionPlan.objects.all().order_by('plan_id')
    return render(request, 'manage_subscriptions.html', {'subscriptions': subscriptions})

from django.shortcuts import render
from .models import Subscription

def subscribed_users(request):
    # Fetch all subscriptions for the subscribed users table
    subscriptions = Subscription.objects.select_related('plan_id').order_by('plan_id')
    return render(request, 'subscribed_users.html', {'subscriptions': subscriptions})

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import SubscriptionPlan

@csrf_exempt
def update_subscription_plan(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            plan_id = data.get('plan_id')
            plan_name = data.get('plan_name')
            price = data.get('price')
            duration_months = data.get('duration_months')

            # Update the subscription plan
            plan = SubscriptionPlan.objects.get(pk=plan_id)
            plan.plan_name = plan_name
            plan.price = price
            plan.duration_months = duration_months
            plan.save()

            return JsonResponse({'success': True})

        except SubscriptionPlan.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Plan not found'}, status=404)

        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
    return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)

def fetch_subscriptions(request):
    free_plan_id = 1
    standard_plan_id = 2  # Replace with your actual logic
    premium_plan_id = 3   # Replace with your actual logic

    free_plan = SubscriptionPlan.objects.filter(plan_id=free_plan_id).first()
    standard_plan = SubscriptionPlan.objects.filter(plan_id=standard_plan_id).first()
    premium_plan = SubscriptionPlan.objects.filter(plan_id=premium_plan_id).first()

    return JsonResponse({
        'standard_price': standard_plan.price if standard_plan else 0,
        'standard_plan_name': standard_plan.plan_name if standard_plan else '',
        'premium_price': premium_plan.price if premium_plan else 0,
        'premium_plan_name': premium_plan.plan_name if premium_plan else '',
        'free_price': free_plan.price if free_plan else 0,
        'free_plan_name': free_plan.plan_name if free_plan else '',
    })


def subscribe_free_trial(request):
    user = request.user  
    plan = SubscriptionPlan.objects.get(plan_name='Free Trial')
    #free_trial_used = Subscription.objects.filter(user_id=user, plan_id_id=plan, status='inactive').exists() 
    #currently_subscribed = Subscription.objects.filter(user_id=user, status='active').exclude(plan_id=plan).exists()
    current_subscription = Subscription.objects.filter(user_id=user, status='active').first()
    free_trial_used = Subscription.objects.filter(user_id=user, plan_id=plan, status='inactive').exists()

    if current_subscription:
        messages.error(request, "You are currently subscribed to another plan and cannot subscribe to the free trial.")
        return redirect('subscribe')
    if not user.is_subscribed and not free_trial_used: 
        Subscription.objects.filter(user_id=user).delete()
        subscription = Subscription.objects.create(
            start_date=datetime.now().date(),
            end_date=(datetime.now() + timedelta(days=30)).date(), 
            user_id=user,
            plan_id=plan, 
            status='active'
        )

       
        user.is_subscribed = True
        user.subscription_status = 'free_trial'
        user.sub_id = subscription.sub_id
        user.save()

       
        messages.success(request, 'You are now subscribed to the Free Trial plan 30days is crazyyyyy.')

          
        return redirect('subscribe') 
    
    else:
         
        if free_trial_used:
       
         messages.error(request, "You have already used the free trial and cannot subscribe again.")
         return redirect('subscribe')

def deactivate_subscription(request):
    

     if request.method == 'POST':
        #subscription_id = request.POST.get('subscription_id')
        user_idd = request.POST.get('user_id')
        try:
            #subscription = Subscription.objects.get(sub_id=subscription_id)
            subscription = Subscription.objects.get(user_id=user_idd)
            subscription.status = 'inactive'  
            
            subscription.save()
            user = subscription.user_id
            user.is_subscribed = False
           
            user.save()
            return JsonResponse({'success': True})
        except Subscription.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Subscription not found'})
     return JsonResponse({'success': False, 'error': 'Invalid request'})